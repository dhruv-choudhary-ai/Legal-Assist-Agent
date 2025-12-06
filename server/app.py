from flask import Flask, request, jsonify, send_file, Response, stream_with_context
import requests
from flask_cors import CORS
from docx import Document
import mammoth
import psycopg2
from dotenv import load_dotenv
import os   
import sys
import logging
import uuid
import json
from pathlib import Path
from flask_bcrypt import Bcrypt
from flask_jwt_extended import JWTManager, create_access_token, jwt_required, get_jwt_identity
from datetime import timedelta
import re

# Setup logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

current_dir = os.path.dirname(os.path.abspath(__file__))
parent_dir = os.path.dirname(current_dir)
sys.path.append(parent_dir)

# Load environment variables
load_dotenv()

# Import AI services
from ai.azure_openai_service import ai_service
from ai.conversation_manager import conversation_manager
from ai.config import AIConfig
from ai.rag_pipeline import rag_pipeline
from ai.vectordb_manager import vector_db
from ai.document_processor import doc_processor
from ai.template_manager_v2 import get_template_manager

app = Flask(__name__)

# Import and register template assembly API blueprint
try:
    from api.template_routes import template_api
    app.register_blueprint(template_api)
    logger.info("✅ Template Assembly API routes registered")
except Exception as e:
    logger.warning(f"⚠️ Could not register template API routes: {e}")

# Configure CORS to allow requests from React frontend
cors = CORS(app, resources={
    r"/api/*": {
        "origins": ["http://localhost:3000", "http://127.0.0.1:3000"],
        "methods": ["GET", "POST", "PUT", "DELETE", "OPTIONS"],
        "allow_headers": ["Content-Type", "Authorization"],
        "supports_credentials": True
    }
})

# Initialize authentication extensions
bcrypt = Bcrypt(app)
app.config['JWT_SECRET_KEY'] = os.getenv('JWT_SECRET_KEY', 'dev-secret-key-change-in-production')
app.config['JWT_ACCESS_TOKEN_EXPIRES'] = timedelta(hours=int(os.getenv('JWT_ACCESS_TOKEN_EXPIRES', 1)))
jwt = JWTManager(app)

# Compatibility decorator: token_required
# Some endpoints in this codebase expect a @token_required decorator that injects a 'current_user' dict.
# This decorator verifies the JWT, looks up the user in the DB and passes current_user as the first arg.
def token_required(fn):
    from functools import wraps
    from flask_jwt_extended import verify_jwt_in_request, get_jwt_identity

    @wraps(fn)
    def wrapper(*args, **kwargs):
        try:
            # Ensure a valid JWT is present in the request
            verify_jwt_in_request()
            user_identity = get_jwt_identity()

            # Convert identity to int if possible
            try:
                user_id = int(user_identity)
            except Exception:
                user_id = user_identity

            # Check if database is available
            if db is None:
                logger.error("Database not available for authentication")
                return jsonify({'error': 'Service temporarily unavailable'}), 503

            # Fetch user details from the database
            cur = db.cursor()
            cur.execute(
                "SELECT user_id, email, full_name, is_active FROM users WHERE user_id = %s AND is_active = TRUE",
                (user_id,)
            )
            user = cur.fetchone()
            cur.close()

            if not user:
                return jsonify({'error': 'User not found or inactive'}), 404

            current_user = {
                'id': user[0],
                'email': user[1],
                'full_name': user[2],
                'is_active': user[3]
            }

            # Call the original function, injecting current_user as the first argument
            return fn(current_user, *args, **kwargs)

        except Exception as e:
            logger.warning(f"Unauthorized access or invalid token: {e}")
            return jsonify({'error': 'Authorization required'}), 401

    return wrapper

# JWT error handlers
@jwt.expired_token_loader
def expired_token_callback(jwt_header, jwt_payload):
    logger.warning(f"⚠️ Expired token attempt")
    return jsonify({'error': 'Token has expired'}), 401

@jwt.invalid_token_loader
def invalid_token_callback(error):
    logger.warning(f"⚠️ Invalid token: {error}")
    return jsonify({'error': 'Invalid token'}), 401

@jwt.unauthorized_loader
def missing_token_callback(error):
    logger.warning(f"⚠️ Missing token: {error}")
    return jsonify({'error': 'Authorization token is missing'}), 401

@jwt.revoked_token_loader
def revoked_token_callback(jwt_header, jwt_payload):
    logger.warning(f"⚠️ Revoked token attempt")
    return jsonify({'error': 'Token has been revoked'}), 401

# Log AI configuration on startup
logger.info("="*60)
logger.info("🚀 Legal Documentation Assistant - AI Backend Starting")
logger.info("="*60)
if AIConfig.validate():
    logger.info("✅ Azure OpenAI Configuration Valid")
    logger.info(f"📊 Configuration: {AIConfig.get_summary()}")
else:
    logger.error("❌ Azure OpenAI not configured - Please set up .env file")
logger.info("="*60)

# Database connection with error handling
try:
    db = psycopg2.connect(
        database=os.getenv('DATABASE_NAME'), 
        user=os.getenv('DATABASE_USER'),
        password=os.getenv('PASSWORD'), 
        host=os.getenv('DATABASE_HOST'), 
        port=os.getenv('DATABASE_PORT'), 
        sslmode='require',  # SSL mode for Render.com
        connect_timeout=10,  # Connection timeout in seconds
        keepalives=1, 
        keepalives_idle=30,
        keepalives_interval=10, 
        keepalives_count=5
    )
    logger.info("✅ Database connected successfully")
except psycopg2.OperationalError as e:
    logger.error(f"❌ Database connection failed: {e}")
    logger.warning("⚠️ Server will start without database (some features may not work)")
    db = None
except Exception as e:
    logger.error(f"❌ Unexpected database error: {e}")
    db = None

# API Routes


@app.route('/api/services', methods=["GET"])
def services():
    if db is None:
        return jsonify({'error': 'Database not available'}), 503
    
    cur = db.cursor()
    cur.execute('SELECT * FROM services')
    row_headers = [x[0] for x in cur.description]
    rv = cur.fetchall()
    json_data = []
    for result in rv:
        json_data.append(dict(zip(row_headers, result)))
    cur.close()
    print(json_data)
    return jsonify(json_data)

# Get forms of a particular service


@app.route('/api/forms', methods=["GET"])
def get_forms():
    # Send json object {"service_id": "..."}
    if db is None:
        return jsonify({'error': 'Database not available'}), 503
    
    Service = request.args.get('service_id')
    print(type(Service))
    print(Service)
    cur = db.cursor()
    cur.execute(
        "SELECT services.service_id, services.service_name, forms.form_id, forms.form_name, forms.form_link FROM services INNER JOIN forms ON services.service_id = forms.service_id WHERE forms.service_id = %s;", [Service])
    row_headers = [x[0] for x in cur.description]
    rv = cur.fetchall()
    json_data = []
    for result in rv:
        json_data.append(dict(zip(row_headers, result)))
    cur.close()
    print(json_data)
    return jsonify(json_data)

# Get all queries for a form


@app.route('/api/form-details', methods=["GET"])
def get_form_details():
    # Send json object {"form_id":"..."}
    if db is None:
        return jsonify({'error': 'Database not available'}), 503
    
    form_id = request.args.get('form_id')
    print(form_id)
    cur = db.cursor()
    cur.execute("SELECT * FROM forms WHERE form_id = %s;", form_id)
    row_headers = [x[0] for x in cur.description]
    rv = cur.fetchall()
    json_data = []
    for result in rv:
        json_data.append(dict(zip(row_headers, result)))
    
    cur.execute("SELECT * FROM ques_categories WHERE id IN (SELECT DISTINCT(category_id) FROM input_ques WHERE ques_id IN (SELECT form_query_id FROM form_queries WHERE form_id = %s));", [form_id])
    row_headers = [x[0] for x in cur.description]
    rv = cur.fetchall()
    for result in rv:
        json_data.append(dict(zip(row_headers, result)))
    cur.execute("SELECT * FROM input_ques WHERE ques_id IN (SELECT form_query_id FROM form_queries WHERE form_id = %s);", [form_id])
    row_headers = [x[0] for x in cur.description]
    rv = cur.fetchall()
    for result in rv:
        json_data.append(dict(zip(row_headers, result)))
    cur.close()
    return jsonify(json_data)


# Return the contents of final doc
@app.route('/api/final-content', methods=["POST"])
def final_content():
    if db is None:
        return jsonify({'error': 'Database not available'}), 503
    
    form_details = request.json                         # Under Progress
    form_id = form_details["form_id"]
    # print(type(form_details))
    print(form_id)
    cur = db.cursor()
    cur.execute("SELECT form_link FROM forms where form_id = %s;", [form_id])
    row_headers = [x[0] for x in cur.description]
    rv = cur.fetchall()
    json_data = []
    for result in rv:
        json_data.append(dict(zip(row_headers, result)))
    cur.close()
    print(json_data[0]["form_link"])
    response = requests.get(json_data[0]["form_link"])
    directory = './docs'
    
    if not os.path.exists(directory):
        os.makedirs(directory)

    file_path = './docs/localfile.docx'

    with open(file_path, 'wb') as f:
        f.write(response.content)
        
    doc = Document('./docs/localfile.docx')
    test = list([int(x) for x in form_details.keys() if x.isdigit()])
    
    test.sort(reverse=True)
    print(test)
    for key in test:
        old = '#'+str(key)
        new = str(form_details[str(key)])
        
        for p in doc.paragraphs:
            if old in p.text:
                # print(old)
                inline = p.runs
                for i in range(len(inline)):
                    if old in inline[i].text:
                        # print(old)
                        res = inline[i].text.replace(old, new)
                        inline[i].text = res
    doc.save("./docs/Output2.docx")
    
    f = open('./docs/Output2.docx', 'rb')
    
    docx_content = mammoth.convert_to_html(f)
    # print(docx_content.value)
    # docx_content.close()

    # fullText = []
    # for para in doc.paragraphs:
    #     fullText.append(para.text)
    # fullText = '\n'.join(fullText)
    # print(fullText)
    return jsonify({'content': docx_content.value})
 
# Return the final doc


@app.route('/api/final-form', methods=["POST"])
def final_form():
    contents = request.get_json()
    print(contents)
    with open('docs/Output2.docx', 'w') as file:
        file.write(contents)
    # doc = Document('docs/localfile.docx')
    # for key, value in form_details.items():
    #     for paragraph in doc.paragraphs:
    #         paragraph.text = paragraph.text.replace(
    #             "#"+str(key)+'#', str(value))

    # doc.save("docs/Output2.docx")
    return send_file('./docs/Output2.docx', as_attachment=True)

@app.route('/api/chat', methods=['POST'])
def chat():
    """
    Standard Chat endpoint with Azure OpenAI (without RAG)
    For general legal questions without knowledge base retrieval
    """
    try:
        user_input = request.json
        user_message = user_input.get('user_chat', '')
        session_id = user_input.get('session_id', str(uuid.uuid4()))
        
        if not user_message:
            return jsonify({'error': 'No message provided'}), 400
        
        if not AIConfig.validate():
            return jsonify({'error': 'AI service not configured'}), 503
        
        logger.info(f"💬 Standard Chat | Session: {session_id} | Message: {user_message[:50]}...")
        
        # Get conversation history
        history = conversation_manager.get_history(session_id, max_messages=10)
        
        # Add user message to history
        conversation_manager.add_message(session_id, 'user', user_message)
        
        # Get AI response without RAG
        response = ai_service.legal_chat(
            user_message=user_message,
            conversation_history=history,
            stream=False
        )
        
        # Add assistant response to history
        conversation_manager.add_message(session_id, 'assistant', response)
        
        return jsonify({
            'aiMessage': response,
            'session_id': session_id,
            'model': 'azure-openai-gpt-4o',
            'usage_stats': ai_service.get_usage_stats()
        })
    
    except Exception as e:
        logger.error(f"❌ Chat error: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/document/analyze', methods=['POST'])
def analyze_document():
    """
    NEW: Analyze legal document using AI
    """
    try:
        data = request.json
        document_type = data.get('document_type', 'legal document')
        document_content = data.get('document_content', '')
        
        if not document_content or not AIConfig.validate():
            return jsonify({'error': 'Invalid request'}), 400
        
        logger.info(f"📄 Analyzing {document_type}")
        
        analysis = ai_service.analyze_document(
            document_type=document_type,
            document_content=document_content,
            stream=False
        )
        
        return jsonify({
            'analysis': analysis,
            'document_type': document_type,
            'model': 'azure-openai-gpt-4o-mini'
        })
    
    except Exception as e:
        logger.error(f"❌ Analysis error: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/document/compare', methods=['POST'])
def compare_documents():
    """
    NEW: Compare two contracts/documents
    """
    try:
        data = request.json
        contract1 = data.get('contract1', '')
        contract2 = data.get('contract2', '')
        
        if not contract1 or not contract2 or not AIConfig.validate():
            return jsonify({'error': 'Invalid request'}), 400
        
        logger.info("📊 Comparing contracts")
        
        comparison = ai_service.compare_contracts(
            contract1=contract1,
            contract2=contract2,
            stream=False
        )
        
        return jsonify({
            'comparison': comparison,
            'model': 'azure-openai-gpt-4o-mini'
        })
    
    except Exception as e:
        logger.error(f"❌ Comparison error: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/form/assist', methods=['POST'])
def assist_form():
    """
    NEW: Assist with form field filling
    """
    try:
        data = request.json
        form_name = data.get('form_name', '')
        field_name = data.get('field_name', '')
        field_description = data.get('field_description', '')
        
        if not all([form_name, field_name, field_description]) or not AIConfig.validate():
            return jsonify({'error': 'Invalid request'}), 400
        
        logger.info(f"📝 Form assist: {form_name} - {field_name}")
        
        assistance = ai_service.assist_form_filling(
            form_name=form_name,
            field_name=field_name,
            field_description=field_description
        )
        
        return jsonify({
            'assistance': assistance,
            'field_name': field_name
        })
    
    except Exception as e:
        logger.error(f"❌ Form assist error: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/legal/question', methods=['POST'])
def answer_question():
    """
    NEW: Answer legal questions with context
    """
    try:
        data = request.json
        question = data.get('question', '')
        document_context = data.get('document_context', '')
        
        if not question or not AIConfig.validate():
            return jsonify({'error': 'Invalid request'}), 400
        
        logger.info(f"❓ Legal question: {question[:50]}...")
        
        answer = ai_service.answer_legal_question(
            question=question,
            document_context=document_context,
            stream=False
        )
        
        return jsonify({
            'answer': answer,
            'question': question
        })
    
    except Exception as e:
        logger.error(f"❌ Question error: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/conversation/clear', methods=['POST'])
def clear_conversation():
    """
    NEW: Clear conversation history
    """
    try:
        data = request.json
        session_id = data.get('session_id', '')
        
        if not session_id:
            return jsonify({'error': 'Session ID required'}), 400
        
        conversation_manager.clear_session(session_id)
        
        return jsonify({
            'message': 'Conversation cleared',
            'session_id': session_id
        })
    
    except Exception as e:
        logger.error(f"❌ Clear error: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/admin/stats', methods=['GET'])
def get_stats():
    """
    NEW: Get AI usage statistics (admin endpoint)
    """
    try:
        return jsonify({
            'ai_usage': ai_service.get_usage_stats(),
            'active_sessions': len(conversation_manager.get_all_sessions()),
            'config': AIConfig.get_summary(),
            'rag_stats': rag_pipeline.get_stats()
        })
    
    except Exception as e:
        logger.error(f"❌ Stats error: {str(e)}")
        return jsonify({'error': str(e)}), 500


# ===================================
# RAG (RETRIEVAL AUGMENTED GENERATION) ENDPOINTS
# ===================================

@app.route('/api/chat/rag', methods=['POST'])
def chat_with_rag():
    """
    Unified Chat with RAG (Retrieval Augmented Generation)
    Like Harvey.ai - Always uses both GPT-4o-mini AND Legal Knowledge Base
    
    Flow:
    1. User sends message
    2. System searches legal knowledge base for relevant documents
    3. GPT-4o-mini generates response using both:
       - Retrieved legal documents (RAG)
       - Its own legal expertise
    4. Response includes citations from knowledge base
    """
    try:
        data = request.json
        user_message = data.get('user_chat', '')
        session_id = data.get('session_id', str(uuid.uuid4()))
        n_results = data.get('n_results', 5)
        
        if not user_message:
            return jsonify({'error': 'No message provided'}), 400
        
        logger.info(f"� Unified AI Chat | Session: {session_id} | Query: {user_message[:50]}...")
        
        # Get conversation history for context
        history = conversation_manager.get_history(session_id, max_messages=10)
        
        # Add user message to history
        conversation_manager.add_message(session_id, 'user', user_message)
        
        # Query with RAG - This ALWAYS uses:
        # 1. Legal Knowledge Base (retrieves relevant documents)
        # 2. GPT-4o-mini (generates intelligent response)
        # 3. Combines both for best results (like Harvey.ai)
        result = rag_pipeline.query_with_rag(
            user_query=user_message,
            conversation_history=history,
            n_results=n_results,
            stream=False,
            include_citations=True
        )
        
        # Add assistant response to history
        if isinstance(result, dict) and 'response' in result:
            conversation_manager.add_message(session_id, 'assistant', result['response'])
            
            return jsonify({
                'success': True,
                'result': {
                    'response': result['response'],
                    'sources': result.get('sources', []),
                    'used_rag': result.get('used_rag', True)
                },
                'session_id': session_id,
                'model': 'gpt-4o-mini + legal-knowledge-base',
                'message': 'Response generated using AI + Legal Knowledge Base'
            })
        else:
            # Fallback if RAG returns just string
            conversation_manager.add_message(session_id, 'assistant', result)
            
            return jsonify({
                'success': True,
                'result': {
                    'response': result,
                    'sources': [],
                    'used_rag': False
                },
                'session_id': session_id,
                'model': 'gpt-4o-mini',
                'message': 'Response generated using AI only (no relevant legal documents found)'
            })
    
    except Exception as e:
        logger.error(f"❌ RAG chat error: {str(e)}")
        return jsonify({
            'error': 'Failed to process your request. Please try again.',
            'details': str(e) if app.debug else None
        }), 500


@app.route('/api/chat/document-query', methods=['POST'])
def document_context_query():
    """
    Handle queries about the current document with full context
    Allows users to ask for summaries, explanations, or modifications
    """
    try:
        data = request.json
        user_query = data.get('user_query', '')
        document_content = data.get('document_content', '')
        document_type = data.get('document_type', 'Legal Document')
        query_type = data.get('query_type', 'general')  # summary, refinement, or general
        session_id = data.get('session_id', str(uuid.uuid4()))
        
        if not user_query:
            return jsonify({'error': 'Query required'}), 400
        
        if not AIConfig.validate():
            return jsonify({'error': 'AI service not configured'}), 503
        
        logger.info(f"📄 Document query | Type: {query_type} | Query: {user_query[:100]}...")
        
        # Create context-aware prompt based on query type
        if query_type == 'summary':
            system_prompt = f"""You are a legal assistant analyzing a {document_type}. 
Provide a clear, concise summary of the document focusing on:
1. Main parties involved
2. Key terms and conditions
3. Important dates and amounts
4. Any critical clauses or obligations

Be professional but conversational."""

            user_prompt = f"""Document Type: {document_type}

Document Content:
{document_content[:3000]}

User Request: {user_query}

Provide a helpful summary or explanation:"""

        elif query_type == 'refinement':
            system_prompt = f"""You are a legal document editor helping refine a {document_type}.
Suggest specific, actionable changes based on the user's request.
Be clear about what to change and why."""

            user_prompt = f"""Document Type: {document_type}

Current Document:
{document_content[:3000]}

User's Change Request: {user_query}

Suggest specific changes:"""

        else:
            system_prompt = f"""You are a legal assistant helping with a {document_type}.
Answer the user's question about the document clearly and helpfully."""

            user_prompt = f"""Document Type: {document_type}

Document Content:
{document_content[:3000]}

User Question: {user_query}

Your Answer:"""

        # Call GPT with document context
        response = ai_service.client.chat.completions.create(
            model=AIConfig.AZURE_OPENAI_CHAT_DEPLOYMENT,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.7,
            max_tokens=1000
        )
        
        assistant_response = response.choices[0].message.content.strip()
        
        logger.info(f"✅ Document query completed")
        
        return jsonify({
            'success': True,
            'response': assistant_response,
            'query_type': query_type,
            'session_id': session_id
        })
    
    except Exception as e:
        logger.error(f"❌ Document query error: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@app.route('/api/knowledge/search', methods=['POST'])
def search_knowledge():
    """
    NEW: Search knowledge base
    """
    try:
        data = request.json
        query = data.get('query', '')
        n_results = data.get('n_results', 10)
        filters = data.get('filters', None)
        
        if not query:
            return jsonify({'error': 'Query required'}), 400
        
        logger.info(f"🔍 Knowledge search: {query[:50]}...")
        
        results = rag_pipeline.search_knowledge_base(
            query=query,
            n_results=n_results,
            filters=filters
        )
        
        return jsonify({
            'success': True,
            'query': query,
            'results': results,
            'count': len(results)
        })
    
    except Exception as e:
        logger.error(f"❌ Search error: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/knowledge/add', methods=['POST'])
def add_to_knowledge():
    """
    NEW: Add document to knowledge base
    """
    try:
        # Check if file upload or path provided
        if 'file' in request.files:
            file = request.files['file']
            if file.filename == '':
                return jsonify({'error': 'No file selected'}), 400
            
            # Save temporarily
            import tempfile
            with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(file.filename)[1]) as tmp:
                file.save(tmp.name)
                tmp_path = tmp.name
            
            document_type = request.form.get('document_type', 'legal document')
            
            result = rag_pipeline.add_document_to_knowledge_base(
                file_path=tmp_path,
                document_type=document_type
            )
            
            # Clean up
            os.unlink(tmp_path)
        
        else:
            # Path-based upload
            data = request.json
            file_path = data.get('file_path', '')
            document_type = data.get('document_type', 'legal document')
            metadata = data.get('metadata', {})
            
            if not file_path:
                return jsonify({'error': 'File path required'}), 400
            
            result = rag_pipeline.add_document_to_knowledge_base(
                file_path=file_path,
                document_type=document_type,
                metadata=metadata
            )
        
        return jsonify(result)
    
    except Exception as e:
        logger.error(f"❌ Add document error: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/knowledge/populate', methods=['POST'])
def populate_knowledge():
    """
    NEW: Populate knowledge base from directory (admin)
    """
    try:
        data = request.json
        directory = data.get('directory', '')
        recursive = data.get('recursive', True)
        
        if not directory:
            return jsonify({'error': 'Directory path required'}), 400
        
        if not os.path.exists(directory):
            return jsonify({'error': 'Directory not found'}), 404
        
        logger.info(f"📚 Populating knowledge base from {directory}")
        
        result = rag_pipeline.populate_knowledge_base(
            directory=directory,
            recursive=recursive
        )
        
        return jsonify(result)
    
    except Exception as e:
        logger.error(f"❌ Populate error: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/knowledge/stats', methods=['GET'])
def knowledge_stats():
    """
    NEW: Get knowledge base statistics
    """
    try:
        stats = rag_pipeline.get_stats()
        return jsonify({
            'success': True,
            'stats': stats
        })
    
    except Exception as e:
        logger.error(f"❌ Stats error: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/knowledge/clear', methods=['POST'])
def clear_knowledge():
    """
    NEW: Clear knowledge base (admin)
    """
    try:
        confirm = request.json.get('confirm', False)
        
        if not confirm:
            return jsonify({'error': 'Confirmation required'}), 400
        
        success = vector_db.delete_collection()
        
        return jsonify({
            'success': success,
            'message': 'Knowledge base cleared' if success else 'Failed to clear'
        })
    
    except Exception as e:
        logger.error(f"❌ Clear error: {str(e)}")
        return jsonify({'error': str(e)}), 500


# ============================================
# NATURAL LANGUAGE DOCUMENT GENERATION ENDPOINTS
# Harvey.ai-style document creation
# ============================================

@app.route('/api/document/generate-from-nl', methods=['POST'])
def generate_document_from_natural_language():
    """
    Generate legal document from natural language description
    
    Example:
    "I need a lease agreement for my office in Mumbai for 11 months at ₹50,000/month"
    
    Flow:
    1. Extract intent (document type, parties, terms)
    2. Generate complete document using AI
    3. Return document + extracted fields + missing fields
    """
    try:
        data = request.json
        user_description = data.get('description', '')
        
        if not user_description:
            return jsonify({'error': 'Description required'}), 400
        
        if not AIConfig.validate():
            return jsonify({'error': 'AI service not configured'}), 503
        
        logger.info(f"📝 Natural language document generation | Description: {user_description[:100]}...")
        
        # Step 1: Extract intent
        intent = ai_service.extract_document_intent(user_description)
        
        if intent.get('confidence', 0) < 0.5:
            return jsonify({
                'success': False,
                'message': 'Could not understand document type. Please be more specific.',
                'intent': intent
            }), 400
        
        # Step 2: Generate document
        document_content = ai_service.generate_document_from_description(
            user_description=user_description,
            extracted_intent=intent,
            stream=False
        )
        
        # Step 3: Check if there are missing fields
        missing_fields = intent.get('missing_fields', [])
        needs_more_info = len(missing_fields) > 0
        
        # Generate follow-up question if needed
        next_question = None
        if needs_more_info:
            next_question = ai_service.ask_for_missing_information(
                document_type=intent.get('document_type'),
                current_document=document_content,
                extracted_fields=intent.get('extracted_fields', {}),
                missing_fields=missing_fields
            )
        
        return jsonify({
            'success': True,
            'document': document_content,
            'document_type': intent.get('document_type'),
            'category': intent.get('category'),
            'extracted_fields': intent.get('extracted_fields', {}),
            'missing_fields': missing_fields,
            'needs_more_info': needs_more_info,
            'next_question': next_question,
            'confidence': intent.get('confidence'),
            'ready_for_use': not needs_more_info
        })
    
    except Exception as e:
        logger.error(f"❌ Natural language generation error: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/document/smart-extract', methods=['POST'])
def smart_extract_fields():
    """
    Intelligently extract field values from user's natural language prompt
    
    This endpoint uses GPT-4o-mini to extract structured field values from
    the user's initial prompt to minimize redundant questions and improve UX.
    
    Flow:
    1. Receive: user_prompt + template_name
    2. Use GPT-4o-mini to extract field values (minimal tokens ~500)
    3. Return: extracted fields + missing fields
    
    Example:
    Input: "I need a lease agreement for my office in Mumbai. The tenant is John Doe,
            rent is ₹50,000 per month for 11 months starting Jan 1st 2024"
    Output: {
        extracted_fields: {
            property_type: "office",
            property_address: "Mumbai",
            tenant_name: "John Doe",
            rent_amount: "50000",
            lease_duration: "11 months",
            start_date: "2024-01-01"
        },
        missing_fields: ["landlord_name", "security_deposit"],
        confidence: 0.85
    }
    """
    try:
        data = request.json
        user_prompt = data.get('prompt', '')
        template_name = data.get('template_name', '')
        
        if not user_prompt:
            return jsonify({'error': 'User prompt required'}), 400
        
        if not template_name:
            return jsonify({'error': 'Template name required'}), 400
        
        if not AIConfig.validate():
            return jsonify({'error': 'AI service not configured'}), 503
        
        logger.info(f"🧠 Smart field extraction | Template: {template_name} | Prompt: {user_prompt[:100]}...")
        
        # Load template schema from template_manager_v2
        tm = get_template_manager()
        schema = tm.get_template_schema(template_name)
        
        if not schema:
            # Fallback schema if template not found in config
            schema = {
                "fields": ["party1_name", "party2_name", "effective_date"],
                "required": ["party1_name", "party2_name"]
            }
            logger.warning(f"⚠️ Template '{template_name}' not found in config, using fallback schema")
        
        # Create GPT-4o-mini extraction prompt (minimal tokens)
        extraction_prompt = f"""You are a legal document assistant. Extract field values from the user's request.

User Request: "{user_prompt}"

Document Type: {template_name}

Extract values for these fields:
{json.dumps(schema['fields'], indent=2)}

IMPORTANT RULES:
1. Only extract information explicitly mentioned in the user's request
2. If a field is not mentioned, leave it as null
3. Use consistent formatting:
   - Dates: YYYY-MM-DD format
   - Money: numeric value only (no currency symbols)
   - Names: proper capitalization
4. Be precise and conservative - don't guess or infer information

Return ONLY a valid JSON object with this exact structure:
{{
    "extracted_fields": {{
        // Only include fields that were explicitly mentioned
    }},
    "confidence": 0.0-1.0  // How confident are you in the extraction?
}}"""

        # Call GPT-4o-mini for extraction (cheap, fast, ~500 tokens)
        try:
            response = ai_service.client.chat.completions.create(
                model=AIConfig.AZURE_OPENAI_CHAT_DEPLOYMENT,  # Use configured deployment
                messages=[
                    {"role": "system", "content": "You are a precise legal data extraction assistant. Return only valid JSON."},
                    {"role": "user", "content": extraction_prompt}
                ],
                temperature=0.1,  # Low temperature for consistency
                max_tokens=500,   # Minimal token usage
                response_format={"type": "json_object"}  # Ensure JSON output
            )
            
            # Parse GPT response
            extraction_result = json.loads(response.choices[0].message.content)
            extracted_fields = extraction_result.get('extracted_fields', {})
            confidence = extraction_result.get('confidence', 0.5)
            
            # Determine missing required fields (build from schema['fields'])
            required_fields = [
                field_name for field_name, field_config in schema['fields'].items()
                if field_config.get('required', False)
            ]
            
            missing_fields = [
                field for field in required_fields 
                if field not in extracted_fields or extracted_fields.get(field) is None or extracted_fields.get(field) == 'null' or extracted_fields.get(field) == ''
            ]
            
            # Also include optional fields that weren't mentioned
            all_missing = [
                field for field in schema['fields']
                if field not in extracted_fields or extracted_fields.get(field) is None or extracted_fields.get(field) == 'null' or extracted_fields.get(field) == ''
            ]
            
            logger.info(f"✅ Extracted {len(extracted_fields)} fields | Missing {len(missing_fields)} required fields")
            logger.info(f"   Required fields: {required_fields}")
            logger.info(f"   Missing required: {missing_fields}")
            
            return jsonify({
                'success': True,
                'extracted_fields': extracted_fields,
                'missing_fields': missing_fields,
                'all_missing_fields': all_missing,
                'confidence': confidence,
                'needs_more_info': len(missing_fields) > 0,
                'template_name': template_name,
                'tokens_used': response.usage.total_tokens if hasattr(response, 'usage') else None
            })
            
        except json.JSONDecodeError as je:
            logger.error(f"❌ Failed to parse GPT JSON response: {str(je)}")
            return jsonify({'error': 'Invalid extraction response format'}), 500
            
    except Exception as e:
        logger.error(f"❌ Smart extraction error: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/document/generate-from-template', methods=['POST'])
def generate_from_template():
    """
    Generate a filled document from Jinja2 template
    
    Takes template name + extracted field values and produces a filled .docx document
    
    Request:
    {
        "template_name": "Lease Agreement",
        "field_values": {
            "lessor_name": "John Doe",
            "lessee_name": "Jane Smith",
            ...
        }
    }
    
    Returns: Filled document as .docx file or HTML preview
    """
    try:
        data = request.json
        template_name = data.get('template_name', '')
        field_values = data.get('field_values', {})
        return_format = data.get('format', 'html')  # 'html' or 'docx'
        
        if not template_name:
            return jsonify({'error': 'Template name required'}), 400
        
        logger.info(f"📄 Generating document from template: {template_name}")
        
        # Get template manager
        tm = get_template_manager()
        
        # Fill template
        try:
            filled_doc = tm.fill_template(template_name, field_values)
            
            # Save to temporary file
            output_filename = f"generated_{template_name.replace(' ', '_')}_{uuid.uuid4().hex[:8]}.docx"
            output_path = os.path.join('generated_documents', output_filename)
            
            # Ensure directory exists
            os.makedirs('generated_documents', exist_ok=True)
            
            filled_doc.save(output_path)
            logger.info(f"✅ Document saved: {output_path}")
            
            if return_format == 'docx':
                # Return as downloadable file
                return send_file(
                    output_path,
                    mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    as_attachment=True,
                    download_name=output_filename
                )
            else:
                # Convert to HTML for preview
                with open(output_path, 'rb') as f:
                    result = mammoth.convert_to_html(f)
                    html_content = result.value
                
                return jsonify({
                    'success': True,
                    'html_content': html_content,
                    'download_url': f'/api/document/download/{output_filename}',
                    'filename': output_filename
                })
        
        except FileNotFoundError as e:
            return jsonify({'error': f'Template file not found: {str(e)}'}), 404
        except Exception as e:
            logger.error(f"❌ Template filling error: {str(e)}")
            return jsonify({'error': f'Failed to fill template: {str(e)}'}), 500
    
    except Exception as e:
        logger.error(f"❌ Document generation error: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/document/validate', methods=['POST'])
def validate_document():
    """
    Validate and verify legal document
    
    Checks for:
    - Legal compliance
    - Missing clauses
    - Potential issues
    - Enforceability
    
    Returns detailed validation report with suggested corrections
    """
    try:
        data = request.json
        document_content = data.get('content', '')
        document_type = data.get('document_type', 'Legal Document')
        jurisdiction = data.get('jurisdiction', 'India')
        
        if not document_content:
            return jsonify({'error': 'Document content required'}), 400
        
        if not AIConfig.validate():
            return jsonify({'error': 'AI service not configured'}), 503
        
        logger.info(f"🔍 Validating {document_type}")
        
        # Validate document
        validation_result = ai_service.validate_legal_document(
            document_content=document_content,
            document_type=document_type,
            jurisdiction=jurisdiction
        )
        
        # Apply corrections if requested
        apply_corrections = data.get('apply_corrections', False)
        corrected_document = None
        
        if apply_corrections and validation_result.get('suggested_corrections'):
            # Apply all suggested corrections
            corrected_document = document_content
            for correction in validation_result.get('suggested_corrections', []):
                if 'original' in correction and 'corrected' in correction:
                    corrected_document = corrected_document.replace(
                        correction['original'],
                        correction['corrected']
                    )
        
        return jsonify({
            'success': True,
            'validation': validation_result,
            'corrected_document': corrected_document,
            'overall_status': validation_result.get('overall_status'),
            'compliance_score': validation_result.get('compliance_score'),
            'ready_for_execution': validation_result.get('ready_for_execution'),
            'critical_issues': [
                issue for issue in validation_result.get('issues', [])
                if issue.get('severity') in ['critical', 'high']
            ]
        })
    
    except Exception as e:
        logger.error(f"❌ Validation error: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/document/refine', methods=['POST'])
def refine_document():
    """
    Refine document based on additional user input
    
    Used in conversational document creation:
    User: "I need a lease"
    AI: Generated lease
    User: "Make the security deposit 3 months rent"
    AI: Updates document
    """
    try:
        data = request.json
        current_document = data.get('current_document', '')
        user_instruction = data.get('instruction', '')
        document_type = data.get('document_type', 'Legal Document')
        
        if not current_document or not user_instruction:
            return jsonify({'error': 'Both current document and instruction required'}), 400
        
        if not AIConfig.validate():
            return jsonify({'error': 'AI service not configured'}), 503
        
        logger.info(f"✏️ Refining {document_type} | Instruction: {user_instruction[:50]}...")
        
        # Use AI to refine the document
        system_prompt = f"""You are an expert legal document drafter. You have a {document_type} that needs to be updated based on user instructions.

RULES:
1. Keep all existing content unless explicitly changed
2. Maintain legal language and structure
3. Ensure changes are legally sound
4. Preserve formatting and clause numbering
5. Only modify what the user asks to change"""

        user_prompt = f"""CURRENT DOCUMENT:
{current_document}

USER INSTRUCTION:
{user_instruction}

Update the document according to the instruction. Return the complete updated document."""

        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ]
        
        refined_document = ai_service.chat_completion(messages, temperature=0.5, max_tokens=4000)
        
        return jsonify({
            'success': True,
            'refined_document': refined_document,
            'instruction_applied': user_instruction
        })
    
    except Exception as e:
        logger.error(f"❌ Refinement error: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/document/export', methods=['POST'])
def export_document():
    """
    Export document in multiple formats (DOCX, PDF, HTML, TXT)
    """
    try:
        data = request.json
        document_content = data.get('content', '')
        format_type = data.get('format', 'docx').lower()
        document_title = data.get('title', 'Legal_Document')
        
        logger.info(f"📥 Export request | Format: {format_type} | Title: {document_title}")
        
        if not document_content:
            return jsonify({'error': 'Document content required'}), 400
        
        # Sanitize filename
        safe_title = re.sub(r'[^\w\s-]', '', document_title).strip().replace(' ', '_')
        
        if format_type == 'docx':
            # Use existing DOCX generation (already implemented)
            from docx import Document
            from docx.shared import Pt, Inches
            import io
            from bs4 import BeautifulSoup
            
            logger.info("📄 Generating DOCX...")
            
            doc = Document()
            
            # Add title
            title = doc.add_heading(document_title, 0)
            
            # Add content (parse HTML if present)
            soup = BeautifulSoup(document_content, 'html.parser')
            text = soup.get_text()
            
            for paragraph in text.split('\n'):
                if paragraph.strip():
                    doc.add_paragraph(paragraph.strip())
            
            # Save to bytes
            file_stream = io.BytesIO()
            doc.save(file_stream)
            file_stream.seek(0)
            
            logger.info(f"✅ DOCX generated successfully | Size: {file_stream.getbuffer().nbytes} bytes")
            
            return send_file(
                file_stream,
                as_attachment=True,
                download_name=f'{safe_title}.docx',
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
        
        elif format_type == 'pdf':
            # PDF export using reportlab
            try:
                from reportlab.lib.pagesizes import A4
                from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                from reportlab.lib.units import inch
                from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER
                from bs4 import BeautifulSoup
                import io
                
                logger.info("📄 Generating PDF...")
                
                # Create PDF in memory
                buffer = io.BytesIO()
                doc = SimpleDocTemplate(
                    buffer,
                    pagesize=A4,
                    rightMargin=72,
                    leftMargin=72,
                    topMargin=72,
                    bottomMargin=72
                )
                
                # Container for PDF elements
                story = []
                
                # Styles
                styles = getSampleStyleSheet()
                
                # Title style
                title_style = ParagraphStyle(
                    'CustomTitle',
                    parent=styles['Heading1'],
                    fontSize=24,
                    textColor='#1e293b',
                    spaceAfter=30,
                    alignment=TA_CENTER,
                    fontName='Helvetica-Bold'
                )
                
                # Body style
                body_style = ParagraphStyle(
                    'CustomBody',
                    parent=styles['BodyText'],
                    fontSize=12,
                    alignment=TA_JUSTIFY,
                    spaceAfter=12,
                    fontName='Times-Roman',
                    leading=16
                )
                
                # Add title
                title = Paragraph(document_title, title_style)
                story.append(title)
                story.append(Spacer(1, 0.3*inch))
                
                # Parse HTML content
                soup = BeautifulSoup(document_content, 'html.parser')
                text = soup.get_text()
                
                # Add content paragraphs
                for para in text.split('\n'):
                    if para.strip():
                        p = Paragraph(para.strip(), body_style)
                        story.append(p)
                        story.append(Spacer(1, 0.1*inch))
                
                # Build PDF
                doc.build(story)
                
                # Get PDF data
                pdf_data = buffer.getvalue()
                buffer.close()
                
                logger.info(f"✅ PDF generated successfully | Size: {len(pdf_data)} bytes")
                
                return Response(
                    pdf_data,
                    mimetype='application/pdf',
                    headers={'Content-Disposition': f'attachment; filename={safe_title}.pdf'}
                )
                
            except ImportError:
                logger.warning("⚠️ reportlab not installed, falling back to error message")
                return jsonify({
                    'error': 'PDF export requires reportlab. Install with: pip install reportlab',
                    'fallback': 'Please use DOCX export instead'
                }), 501
            except Exception as pdf_error:
                logger.error(f"❌ PDF generation error: {str(pdf_error)}")
                return jsonify({'error': f'PDF generation failed: {str(pdf_error)}'}), 500
        
        elif format_type == 'html':
            # HTML export
            html_content = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{document_title}</title>
    <style>
        body {{ font-family: 'Times New Roman', serif; margin: 2cm; }}
        h1 {{ text-align: center; }}
        p {{ text-align: justify; line-height: 1.6; }}
    </style>
</head>
<body>
    <h1>{document_title}</h1>
    {document_content}
</body>
</html>"""
            
            return Response(
                html_content,
                mimetype='text/html',
                headers={'Content-Disposition': f'attachment; filename={safe_title}.html'}
            )
        
        elif format_type == 'txt':
            # Plain text export
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(document_content, 'html.parser')
            text = soup.get_text()
            
            return Response(
                text,
                mimetype='text/plain',
                headers={'Content-Disposition': f'attachment; filename={safe_title}.txt'}
            )
        
        else:
            return jsonify({'error': 'Unsupported format. Use: docx, pdf, html, txt'}), 400
    
    except Exception as e:
        logger.error(f"❌ Export error: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/document/validate-enhanced', methods=['POST'])
def validate_document_enhanced():
    """
    Enhanced validation using dual-model verification system
    
    Implements Harvey.ai-style multi-layered verification:
    - Citation verification against Indian Code
    - Clause-level risk analysis
    - Dual-model verification (Generator + Verifier)
    - Self-consistency checking
    - Temporal and jurisdictional awareness
    
    Returns comprehensive verification report with scores and recommendations
    """
    try:
        data = request.json
        document_content = data.get('content', '')
        document_type = data.get('document_type', 'Legal Document')
        verification_level = data.get('verification_level', 'comprehensive')  # basic, standard, comprehensive
        
        if not document_content:
            return jsonify({'error': 'Document content required'}), 400
        
        if not AIConfig.validate():
            return jsonify({'error': 'AI service not configured'}), 503
        
        if verification_level not in ['basic', 'standard', 'comprehensive']:
            return jsonify({'error': 'Invalid verification level. Use: basic, standard, or comprehensive'}), 400
        
        logger.info(f"🔬 Enhanced validation: {document_type} ({verification_level})")
        
        # Run enhanced validation with dual-model verification
        validation_result = ai_service.validate_document_with_verifier(
            document_content=document_content,
            document_type=document_type,
            verification_level=verification_level
        )
        
        # Categorize issues by priority
        critical_issues = []
        high_priority = []
        medium_priority = []
        low_priority = []
        
        for issue in validation_result.get('issues', []):
            severity = issue.get('severity', 'low')
            if severity == 'critical':
                critical_issues.append(issue)
            elif severity == 'high':
                high_priority.append(issue)
            elif severity == 'medium':
                medium_priority.append(issue)
            else:
                low_priority.append(issue)
        
        # Build response
        response = {
            'success': True,
            'verification_level': verification_level,
            'overall_score': validation_result.get('overall_score', 0),
            'compliance_score': validation_result.get('compliance_score', 0),
            'ready_for_execution': validation_result.get('ready_for_execution', False),
            'enforceability_rating': validation_result.get('enforceability_rating', 'medium'),
            
            # Citation verification
            'citation_verification': validation_result.get('citation_verification', {}),
            
            # Categorized issues
            'critical_issues': critical_issues,
            'high_priority_issues': high_priority,
            'medium_priority_issues': medium_priority,
            'low_priority_issues': low_priority,
            'total_issues': len(validation_result.get('issues', [])),
            
            # Detailed analysis
            'clause_analysis': validation_result.get('clause_analysis', []),
            'hallucinations_detected': validation_result.get('hallucinations_detected', []),
            'risky_clauses': validation_result.get('risky_clauses', []),
            'missing_clauses': validation_result.get('missing_clauses', []),
            'suggested_corrections': validation_result.get('suggested_corrections', []),
            
            # Temporal and jurisdictional checks
            'temporal_check': validation_result.get('temporal_check', {}),
            'jurisdictional_check': validation_result.get('jurisdictional_check', {}),
            
            # Consistency score (only for comprehensive verification)
            'consistency_score': validation_result.get('consistency_score'),
            
            # Recommendations
            'recommendations': validation_result.get('recommendations', []),
            
            # Action items
            'action_items': {
                'must_fix': critical_issues,
                'should_fix': high_priority,
                'consider_fixing': medium_priority,
                'optional': low_priority
            }
        }
        
        logger.info(f"✅ Enhanced validation complete | Overall: {response['overall_score']}/100 | Compliance: {response['compliance_score']}/100")
        
        return jsonify(response)
    
    except Exception as e:
        logger.error(f"❌ Enhanced validation error: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/document/fix-issue', methods=['POST'])
def fix_single_issue():
    """
    Fix a single validation issue in the document using AI
    
    Request:
    {
        "document_html": "<html>...</html>",
        "issue": {
            "severity": "high",
            "issue_type": "Missing Clause",
            "description": "...",
            "suggestion": "...",
            "location": "..."
        }
    }
    """
    try:
        data = request.json
        document_html = data.get('document_html', '')
        issue = data.get('issue', {})
        
        if not document_html or not issue:
            return jsonify({'error': 'Document and issue required'}), 400
        
        if not AIConfig.validate():
            return jsonify({'error': 'AI service not configured'}), 503
        
        logger.info(f"🔧 Fixing single issue: {issue.get('issue_type', 'Unknown')}")
        
        # Create fix prompt
        fix_prompt = f"""You are a legal document editor. Fix the following issue in the document.

**Issue Details:**
- Type: {issue.get('issue_type', issue.get('issue', 'Unknown'))}
- Severity: {issue.get('severity', 'medium')}
- Description: {issue.get('description', issue.get('issue', ''))}
- Location: {issue.get('location', issue.get('clause_reference', 'Not specified'))}
- Suggestion: {issue.get('suggestion', issue.get('recommendation', ''))}

**Current Document (HTML):**
{document_html}

**Instructions:**
1. Locate the specific issue in the document
2. Apply the suggested fix or make appropriate corrections
3. Maintain all HTML formatting and structure
4. Return ONLY the complete corrected HTML document
5. Do not add explanations or comments

Return the fixed HTML document:"""

        # Call GPT to fix the issue
        response = ai_service.client.chat.completions.create(
            model=AIConfig.AZURE_OPENAI_CHAT_DEPLOYMENT,
            messages=[
                {"role": "system", "content": "You are a legal document editor. Return only the corrected HTML document without any additional text."},
                {"role": "user", "content": fix_prompt}
            ],
            temperature=0.3,
            max_tokens=4000
        )
        
        fixed_document = response.choices[0].message.content.strip()
        
        # Clean up markdown code blocks if present
        if fixed_document.startswith('```html'):
            fixed_document = fixed_document.replace('```html', '').replace('```', '').strip()
        
        logger.info(f"✅ Issue fixed successfully")
        
        return jsonify({
            'success': True,
            'fixed_document': fixed_document,
            'issue_fixed': issue.get('issue_type', issue.get('issue', 'Unknown'))
        })
    
    except Exception as e:
        logger.error(f"❌ Fix issue error: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/document/fix-all-issues', methods=['POST'])
def fix_all_issues():
    """
    Fix all validation issues in the document at once using AI
    
    Request:
    {
        "document_html": "<html>...</html>",
        "issues": [ {...}, {...}, ... ]
    }
    """
    try:
        data = request.json
        document_html = data.get('document_html', '')
        issues = data.get('issues', [])
        
        if not document_html or not issues:
            return jsonify({'error': 'Document and issues required'}), 400
        
        if not AIConfig.validate():
            return jsonify({'error': 'AI service not configured'}), 503
        
        logger.info(f"🔧 Fixing all {len(issues)} issues")
        
        # Create comprehensive fix prompt
        issues_summary = "\n".join([
            f"{i+1}. [{issue.get('severity', 'medium').upper()}] {issue.get('issue_type', issue.get('issue', 'Unknown'))}: {issue.get('description', issue.get('issue', ''))}"
            for i, issue in enumerate(issues)
        ])
        
        suggestions_summary = "\n".join([
            f"{i+1}. {issue.get('suggestion', issue.get('recommendation', 'Fix as needed'))}"
            for i, issue in enumerate(issues)
        ])
        
        fix_prompt = f"""You are a legal document editor. Fix ALL the following issues in the document.

**Issues to Fix ({len(issues)} total):**
{issues_summary}

**Suggested Fixes:**
{suggestions_summary}

**Current Document (HTML):**
{document_html}

**Instructions:**
1. Fix ALL issues listed above
2. Apply all suggested corrections
3. Maintain all HTML formatting and structure
4. Ensure legal compliance and completeness
5. Return ONLY the complete corrected HTML document
6. Do not add explanations or comments

Return the fully corrected HTML document:"""

        # Call GPT to fix all issues
        response = ai_service.client.chat.completions.create(
            model=AIConfig.AZURE_OPENAI_CHAT_DEPLOYMENT,
            messages=[
                {"role": "system", "content": "You are a legal document editor. Return only the fully corrected HTML document without any additional text."},
                {"role": "user", "content": fix_prompt}
            ],
            temperature=0.3,
            max_tokens=6000
        )
        
        fixed_document = response.choices[0].message.content.strip()
        
        # Clean up markdown code blocks if present
        if fixed_document.startswith('```html'):
            fixed_document = fixed_document.replace('```html', '').replace('```', '').strip()
        
        logger.info(f"✅ All {len(issues)} issues fixed successfully")
        
        return jsonify({
            'success': True,
            'fixed_document': fixed_document,
            'issues_fixed': len(issues)
        })
    
    except Exception as e:
        logger.error(f"❌ Fix all issues error: {str(e)}")
        return jsonify({'error': str(e)}), 500


# ============================================
# DOCUMENT ANALYZER ENDPOINTS (Upload & Analysis)
# ============================================

@app.route('/api/document/analyze/upload', methods=['POST'])
def upload_document_for_analysis():
    """
    Upload a legal document for analysis (PDF/DOCX)
    Session-based storage - no persistence
    """
    try:
        from ai.document_analyzer import document_analyzer
        
        if 'file' not in request.files:
            return jsonify({'error': 'No file uploaded'}), 400
        
        file = request.files['file']
        
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        # Check file extension
        allowed_extensions = ['.pdf', '.docx', '.doc']
        file_ext = Path(file.filename).suffix.lower()
        
        if file_ext not in allowed_extensions:
            return jsonify({'error': f'Unsupported file type. Allowed: {", ".join(allowed_extensions)}'}), 400
        
        # Save temporarily
        import tempfile
        with tempfile.NamedTemporaryFile(delete=False, suffix=file_ext) as tmp:
            file.save(tmp.name)
            tmp_path = tmp.name
        
        logger.info(f"📄 Processing uploaded document: {file.filename}")
        
        # Process document
        doc_id = document_analyzer.process_document(tmp_path, file.filename)
        doc_info = document_analyzer.get_document_info(doc_id)
        
        # Clean up temp file
        os.unlink(tmp_path)
        
        return jsonify({
            'success': True,
            'document_id': doc_id,
            'filename': doc_info['filename'],
            'total_chunks': doc_info['total_chunks'],
            'word_count': doc_info['word_count'],
            'char_count': doc_info['char_count'],
            'message': f'Document "{file.filename}" uploaded and processed successfully!'
        })
        
    except Exception as e:
        logger.error(f"❌ Document upload error: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/document/analyze/question', methods=['POST'])
def answer_document_question():
    """
    Answer a question about the uploaded document using RAG
    Uses BGE-M3 embeddings for efficient retrieval
    """
    try:
        from ai.document_analyzer import document_analyzer
        
        data = request.json
        doc_id = data.get('document_id')
        question = data.get('question')
        
        if not doc_id or not question:
            return jsonify({'error': 'document_id and question required'}), 400
        
        logger.info(f"❓ Question about doc {doc_id}: {question[:50]}...")
        
        result = document_analyzer.answer_question(doc_id, question)
        
        return jsonify({
            'success': True,
            'answer': result['answer'],
            'sources': result['sources'],
            'chunks_used': result['chunks_used']
        })
        
    except Exception as e:
        logger.error(f"❌ Question answering error: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/document/analyze/summarize', methods=['POST'])
def summarize_document():
    """Generate document summary"""
    try:
        from ai.document_analyzer import document_analyzer
        
        data = request.json
        doc_id = data.get('document_id')
        
        if not doc_id:
            return jsonify({'error': 'document_id required'}), 400
        
        logger.info(f"📝 Summarizing document {doc_id}")
        
        result = document_analyzer.summarize_document(doc_id)
        
        return jsonify({
            'success': True,
            'summary': result['summary'],
            'chunks_analyzed': result['chunks_analyzed'],
            'total_chunks': result['total_chunks']
        })
        
    except Exception as e:
        logger.error(f"❌ Summarization error: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/document/analyze/clauses', methods=['POST'])
def extract_document_clauses():
    """Extract key legal clauses from document"""
    try:
        from ai.document_analyzer import document_analyzer
        
        data = request.json
        doc_id = data.get('document_id')
        
        if not doc_id:
            return jsonify({'error': 'document_id required'}), 400
        
        logger.info(f"📋 Extracting clauses from document {doc_id}")
        
        result = document_analyzer.extract_key_clauses(doc_id)
        
        return jsonify({
            'success': True,
            'clauses': result['clauses'],
            'chunks_analyzed': result['chunks_analyzed']
        })
        
    except Exception as e:
        logger.error(f"❌ Clause extraction error: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/document/analyze/risks', methods=['POST'])
def analyze_document_risks():
    """Analyze potential legal risks in document"""
    try:
        from ai.document_analyzer import document_analyzer
        
        data = request.json
        doc_id = data.get('document_id')
        
        if not doc_id:
            return jsonify({'error': 'document_id required'}), 400
        
        logger.info(f"⚠️ Analyzing risks in document {doc_id}")
        
        result = document_analyzer.analyze_risks(doc_id)
        
        return jsonify({
            'success': True,
            'risks': result['risks'],
            'chunks_analyzed': result['chunks_analyzed']
        })
        
    except Exception as e:
        logger.error(f"❌ Risk analysis error: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/document/analyze/clear', methods=['POST'])
def clear_analyzed_document():
    """Clear uploaded document from session"""
    try:
        from ai.document_analyzer import document_analyzer
        
        data = request.json
        doc_id = data.get('document_id')
        
        if not doc_id:
            return jsonify({'error': 'document_id required'}), 400
        
        document_analyzer.clear_document(doc_id)
        
        return jsonify({
            'success': True,
            'message': 'Document cleared from session'
        })
        
    except Exception as e:
        logger.error(f"❌ Clear document error: {str(e)}")
        return jsonify({'error': str(e)}), 500


# ============================================
# LAWYER FEEDBACK SYSTEM ENDPOINTS
# ============================================

@app.route('/api/feedback/submit-correction', methods=['POST'])
@token_required
def submit_feedback_correction(current_user):
    """
    Submit a correction from a lawyer or user
    
    Request body:
    {
        "document_id": 123,
        "original_text": "AI-generated text",
        "corrected_text": "Lawyer-corrected text",
        "correction_type": "legal|grammar|factual",
        "explanation": "Why this was wrong",
        "applicable_law": "Indian Contract Act, Section 10",
        "lawyer_name": "Optional lawyer name",
        "severity": "critical|high|medium|low"
    }
    """
    try:
        from ai.lawyer_feedback import LawyerFeedbackSystem
        
        data = request.json
        
        # Initialize feedback system
        db_config = {
            'host': os.getenv('DB_HOST', 'localhost'),
            'port': int(os.getenv('DB_PORT', 5432)),
            'database': os.getenv('DB_NAME', 'legal_assistant'),
            'user': os.getenv('DB_USER', 'postgres'),
            'password': os.getenv('DB_PASSWORD', '')
        }
        
        feedback_system = LawyerFeedbackSystem(db_config)
        
        # Submit correction
        correction_id = feedback_system.submit_correction(
            document_id=data.get('document_id'),
            user_id=current_user['id'],
            original_text=data.get('original_text'),
            corrected_text=data.get('corrected_text'),
            correction_type=data.get('correction_type', 'general'),
            explanation=data.get('explanation'),
            applicable_law=data.get('applicable_law'),
            lawyer_name=data.get('lawyer_name'),
            severity=data.get('severity', 'medium')
        )
        
        return jsonify({
            'success': True,
            'correction_id': correction_id,
            'message': 'Correction submitted successfully'
        })
    
    except Exception as e:
        logger.error(f"❌ Feedback submission error: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/feedback/rate-suggestion', methods=['POST'])
@token_required
def rate_ai_suggestion_endpoint(current_user):
    """
    Rate an AI suggestion
    
    Request body:
    {
        "suggestion_type": "document_generation",
        "original_query": "User's query",
        "ai_response": "AI's response",
        "rating": 1-5,
        "feedback_text": "Optional feedback",
        "accuracy_score": 0-100
    }
    """
    try:
        from ai.lawyer_feedback import LawyerFeedbackSystem
        
        data = request.json
        
        db_config = {
            'host': os.getenv('DB_HOST', 'localhost'),
            'port': int(os.getenv('DB_PORT', 5432)),
            'database': os.getenv('DB_NAME', 'legal_assistant'),
            'user': os.getenv('DB_USER', 'postgres'),
            'password': os.getenv('DB_PASSWORD', '')
        }
        
        feedback_system = LawyerFeedbackSystem(db_config)
        
        rating_id = feedback_system.rate_ai_suggestion(
            suggestion_type=data.get('suggestion_type'),
            original_query=data.get('original_query'),
            ai_response=data.get('ai_response'),
            rating=data.get('rating'),
            feedback_text=data.get('feedback_text'),
            accuracy_score=data.get('accuracy_score'),
            user_id=current_user['id']
        )
        
        return jsonify({
            'success': True,
            'rating_id': rating_id,
            'message': 'Rating submitted successfully'
        })
    
    except Exception as e:
        logger.error(f"❌ Rating submission error: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/feedback/stats', methods=['GET'])
@token_required
def get_feedback_stats(current_user):
    """Get feedback system statistics"""
    try:
        from ai.lawyer_feedback import LawyerFeedbackSystem
        
        db_config = {
            'host': os.getenv('DB_HOST', 'localhost'),
            'port': int(os.getenv('DB_PORT', 5432)),
            'database': os.getenv('DB_NAME', 'legal_assistant'),
            'user': os.getenv('DB_USER', 'postgres'),
            'password': os.getenv('DB_PASSWORD', '')
        }
        
        feedback_system = LawyerFeedbackSystem(db_config)
        stats = feedback_system.get_feedback_stats()
        
        return jsonify({
            'success': True,
            'stats': stats
        })
    
    except Exception as e:
        logger.error(f"❌ Stats retrieval error: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/ontology/search-clauses', methods=['POST'])
def search_legal_clauses():
    """
    Search legal clauses in ontology
    
    Request body:
    {
        "clause_type": "Optional clause type",
        "legal_area": "Optional legal area",
        "risk_level": "Optional risk level",
        "must_include": true/false,
        "keywords": ["keyword1", "keyword2"]
    }
    """
    try:
        from ai.legal_ontology import legal_ontology
        
        data = request.json or {}
        
        clauses = legal_ontology.search_clauses(
            clause_type=data.get('clause_type'),
            legal_area=data.get('legal_area'),
            risk_level=data.get('risk_level'),
            must_include=data.get('must_include'),
            keywords=data.get('keywords')
        )
        
        return jsonify({
            'success': True,
            'total_results': len(clauses),
            'clauses': [clause.to_dict() for clause in clauses]
        })
    
    except Exception as e:
        logger.error(f"❌ Clause search error: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/ontology/required-clauses/<document_type>', methods=['GET'])
def get_required_clauses(document_type):
    """Get all required clauses for a document type"""
    try:
        from ai.legal_ontology import legal_ontology
        
        clauses = legal_ontology.get_required_clauses(document_type)
        
        return jsonify({
            'success': True,
            'document_type': document_type,
            'total_required': len(clauses),
            'required_clauses': [clause.to_dict() for clause in clauses]
        })
    
    except Exception as e:
        logger.error(f"❌ Required clauses error: {str(e)}")
        return jsonify({'error': str(e)}), 500


# ============================================
# AUTHENTICATION ENDPOINTS
# ============================================

def validate_email(email):
    """Validate email format"""
    pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
    return re.match(pattern, email) is not None

def validate_password(password):
    """Validate password strength (min 8 chars, 1 uppercase, 1 lowercase, 1 number)"""
    if len(password) < 8:
        return False, "Password must be at least 8 characters long"
    if not re.search(r'[A-Z]', password):
        return False, "Password must contain at least one uppercase letter"
    if not re.search(r'[a-z]', password):
        return False, "Password must contain at least one lowercase letter"
    if not re.search(r'\d', password):
        return False, "Password must contain at least one number"
    return True, "Password is valid"

@app.route('/api/auth/signup', methods=['POST'])
def signup():
    """User registration endpoint"""
    try:
        # Check database availability
        if db is None:
            return jsonify({'error': 'Service temporarily unavailable. Please try again later.'}), 503
        
        data = request.json
        email = data.get('email', '').strip().lower()
        password = data.get('password', '')
        full_name = data.get('full_name', '').strip()
        phone = data.get('phone', '').strip()
        
        # Validation
        if not email or not password:
            return jsonify({'error': 'Email and password are required'}), 400
        
        if not validate_email(email):
            return jsonify({'error': 'Invalid email format'}), 400
        
        is_valid, message = validate_password(password)
        if not is_valid:
            return jsonify({'error': message}), 400
        
        # Check if user already exists
        cur = db.cursor()
        cur.execute("SELECT user_id FROM users WHERE email = %s", (email,))
        if cur.fetchone():
            cur.close()
            return jsonify({'error': 'Email already registered'}), 409
        
        # Hash password
        password_hash = bcrypt.generate_password_hash(password).decode('utf-8')
        
        # Insert user
        cur.execute(
            """INSERT INTO users (email, password_hash, full_name, phone) 
               VALUES (%s, %s, %s, %s) RETURNING user_id""",
            (email, password_hash, full_name, phone)
        )
        user_id = cur.fetchone()[0]
        db.commit()
        cur.close()
        
        # Generate JWT token (identity must be string)
        access_token = create_access_token(identity=str(user_id))
        
        logger.info(f"✅ New user registered: {email}")
        logger.info(f"🔑 Generated token for user_id: {user_id}")
        
        return jsonify({
            'success': True,
            'message': 'User registered successfully',
            'token': access_token,
            'user': {
                'user_id': user_id,
                'email': email,
                'full_name': full_name
            }
        }), 201
        
    except Exception as e:
        logger.error(f"❌ Signup error: {str(e)}")
        if db is not None:
            db.rollback()
        return jsonify({'error': 'Registration failed. Please try again.'}), 500

@app.route('/api/auth/login', methods=['POST'])
def login():
    """User login endpoint"""
    try:
        # Check database availability
        if db is None:
            return jsonify({'error': 'Service temporarily unavailable. Please try again later.'}), 503
        
        data = request.json
        email = data.get('email', '').strip().lower()
        password = data.get('password', '')
        
        if not email or not password:
            return jsonify({'error': 'Email and password are required'}), 400
        
        # Get user from database
        cur = db.cursor()
        cur.execute(
            """SELECT user_id, password_hash, full_name, email, is_active 
               FROM users WHERE email = %s""",
            (email,)
        )
        user = cur.fetchone()
        
        if not user:
            cur.close()
            return jsonify({'error': 'Invalid email or password'}), 401
        
        user_id, password_hash, full_name, user_email, is_active = user
        
        if not is_active:
            cur.close()
            return jsonify({'error': 'Account is deactivated'}), 403
        
        # Verify password
        if not bcrypt.check_password_hash(password_hash, password):
            cur.close()
            return jsonify({'error': 'Invalid email or password'}), 401
        
        # Update last login
        cur.execute(
            "UPDATE users SET last_login = CURRENT_TIMESTAMP WHERE user_id = %s",
            (user_id,)
        )
        db.commit()
        cur.close()
        
        # Generate JWT token
        access_token = create_access_token(identity=str(user_id))
        
        logger.info(f"✅ User logged in: {email}")
        logger.info(f"🔑 Generated token for user_id: {user_id}")
        
        return jsonify({
            'success': True,
            'message': 'Login successful',
            'token': access_token,
            'user': {
                'user_id': user_id,
                'email': user_email,
                'full_name': full_name
            }
        }), 200
        
    except Exception as e:
        logger.error(f"❌ Login error: {str(e)}")
        return jsonify({'error': 'Login failed. Please try again.'}), 500

@app.route('/api/auth/verify', methods=['GET'])
@jwt_required()
def verify_token():
    """Verify JWT token and return user info"""
    try:
        # Check database availability
        if db is None:
            return jsonify({'error': 'Service temporarily unavailable. Please try again later.'}), 503
        
        user_id = get_jwt_identity()
        logger.info(f"🔐 Token verification for user_id: {user_id} (type: {type(user_id)})")
        
        # Convert to int if string
        if isinstance(user_id, str):
            try:
                user_id = int(user_id)
            except ValueError:
                logger.error(f"❌ Cannot convert user_id to int: {user_id}")
                return jsonify({'error': 'Invalid user ID format'}), 400
        
        cur = db.cursor()
        cur.execute(
            "SELECT user_id, email, full_name FROM users WHERE user_id = %s AND is_active = TRUE",
            (user_id,)
        )
        user = cur.fetchone()
        cur.close()
        
        if not user:
            logger.warning(f"⚠️ User not found or inactive: {user_id}")
            return jsonify({'error': 'User not found'}), 404
        
        logger.info(f"✅ Token verified for user: {user[1]}")
        return jsonify({
            'success': True,
            'user': {
                'user_id': user[0],
                'email': user[1],
                'full_name': user[2]
            }
        }), 200
        
    except Exception as e:
        logger.error(f"❌ Verification error: {str(e)}")
        return jsonify({'error': 'Token verification failed'}), 401

@app.route('/api/user/profile', methods=['GET'])
@jwt_required()
def get_profile():
    """Get user profile"""
    try:
        if db is None:
            return jsonify({'error': 'Service temporarily unavailable. Please try again later.'}), 503
        
        user_id = get_jwt_identity()
        
        cur = db.cursor()
        cur.execute(
            """SELECT user_id, email, full_name, phone, created_at, last_login 
               FROM users WHERE user_id = %s""",
            (user_id,)
        )
        user = cur.fetchone()
        cur.close()
        
        if not user:
            return jsonify({'error': 'User not found'}), 404
        
        return jsonify({
            'user_id': user[0],
            'email': user[1],
            'full_name': user[2],
            'phone': user[3],
            'created_at': user[4].isoformat() if user[4] else None,
            'last_login': user[5].isoformat() if user[5] else None
        }), 200
        
    except Exception as e:
        logger.error(f"❌ Profile error: {str(e)}")
        return jsonify({'error': 'Failed to fetch profile'}), 500

@app.route('/api/user/documents', methods=['GET'])
@jwt_required()
def get_user_documents():
    """Get user's document history"""
    try:
        if db is None:
            return jsonify({'error': 'Service temporarily unavailable. Please try again later.'}), 503
        
        user_id = get_jwt_identity()
        logger.info(f"📄 Fetching documents for user_id: {user_id}")
        
        cur = db.cursor()
        cur.execute(
            """SELECT doc_id, form_name, created_at, updated_at 
               FROM user_documents 
               WHERE user_id = %s 
               ORDER BY created_at DESC 
               LIMIT 50""",
            (user_id,)
        )
        documents = cur.fetchall()
        cur.close()
        
        doc_list = []
        for doc in documents:
            doc_list.append({
                'doc_id': doc[0],
                'form_name': doc[1],
                'created_at': doc[2].isoformat() if doc[2] else None,
                'updated_at': doc[3].isoformat() if doc[3] else None
            })
        
        logger.info(f"✅ Found {len(doc_list)} documents for user: {user_id}")
        return jsonify({
            'success': True,
            'documents': doc_list,
            'count': len(doc_list)
        }), 200
        
    except Exception as e:
        logger.error(f"❌ Documents error: {str(e)}")
        return jsonify({'error': 'Failed to fetch documents'}), 500


@app.route('/api/document/simple-chat', methods=['POST'])
def simple_document_chat():
    """
    SIMPLE Document Chat - No Chaos
    
    Flow:
    1. User gives prompt → GPT picks template + extracts fields
    2. Missing fields? → Ask naturally
    3. All fields? → Generate document
    4. Add legal clauses via BGE-M3 RAG
    
    Request:
        {
            "message": "create rent agreement for Dhruv...",
            "session_id": "abc123",
            "conversation": [...]  // optional
        }
    
    Response:
        {
            "status": "need_more_info" | "ready",
            "message": "What's the monthly rent?",
            "template": "Lease Agreement",
            "extracted": {"LESSEE_NAME": "Dhruv", ...},
            "missing": ["MONTHLY_RENT", ...],
            "document": "..." (if ready),
            "download_id": "..." (if ready)
        }
    """
    try:
        from ai.simple_assembler import simple_assembler
        import uuid
        
        data = request.json
        user_message = data.get('message', '').strip()
        session_id = data.get('session_id', f"session_{uuid.uuid4()}")
        conversation = data.get('conversation', [])
        
        if not user_message:
            return jsonify({'error': 'message is required'}), 400
        
        logger.info(f"💬 Simple Chat | Session: {session_id[:8]} | Message: {user_message[:80]}...")
        
        # Step 1: Detect template (if not already known)
        template_name = data.get('template')
        if not template_name:
            template_name = simple_assembler.detect_template(user_message)
            if not template_name:
                return jsonify({
                    'status': 'clarify',
                    'message': 'What type of document do you need? (e.g., Lease Agreement, NDA, Legal Notice)',
                    'available_templates': list(simple_assembler.TEMPLATE_CONFIG.keys())
                })
        
        # Step 2: Extract fields from message + conversation
        extraction = simple_assembler.extract_fields(user_message, template_name, conversation)
        
        extracted = extraction['extracted']  # Field names → values (for display)
        extracted_raw = extraction.get('extracted_raw', {})  # Placeholder codes → values (for template)
        missing = extraction['missing']  # Missing placeholder codes
        
        logger.info(f"📊 Template: {template_name} | Extracted: {len(extracted)} | Missing: {len(missing)}")
        
        # Step 3: Check if we have everything
        if not missing:
            # Generate document!
            logger.info("🎉 All fields collected, generating document...")
            
            # Use raw extracted values (with placeholder codes) for template filling
            filled_doc = simple_assembler.fill_template(template_name, extracted_raw)
            
            # Save document
            doc_id = str(uuid.uuid4())[:8]
            output_dir = Path("./generated_documents")
            output_dir.mkdir(exist_ok=True)
            output_path = output_dir / f"{session_id}_{doc_id}.docx"
            filled_doc.save(str(output_path))
            
            # Get text preview
            preview = '\n'.join([p.text for p in filled_doc.paragraphs if p.text.strip()])[:1000]
            
            # Check for extraction artifacts in preview
            if "I told you" in preview or "[" in preview:
                logger.warning("⚠️ Possible extraction artifacts in document")
            
            # Add RAG suggestions
            rag_suggestions = simple_assembler.enhance_with_rag(preview, template_name)
            
            return jsonify({
                'status': 'ready',
                'message': f'✅ Your {template_name} is ready!',
                'template': template_name,
                'extracted': extracted,  # Show user-friendly field names
                'document': preview,
                'rag_suggestions': rag_suggestions,
                'download_id': doc_id,
                'download_url': f'/api/document/download/{doc_id}'
            })
        
        else:
            # Ask for next missing field (pass RAW values with placeholder codes)
            question = simple_assembler.ask_for_missing(missing, template_name, extracted_raw)
            
            return jsonify({
                'status': 'need_more_info',
                'message': question,
                'template': template_name,
                'extracted': extracted,  # Show user-friendly field names
                'missing': missing,
                'progress': {
                    'done': len(extracted),
                    'total': len(extracted) + len(missing),
                    'percent': int((len(extracted) / (len(extracted) + len(missing))) * 100) if (len(extracted) + len(missing)) > 0 else 0
                }
            })
    
    except Exception as e:
        logger.error(f"❌ Simple chat error: {str(e)}", exc_info=True)
        return jsonify({
            'status': 'error',
            'message': 'Sorry, something went wrong. Please try again.',
            'error': str(e)
        }), 500


@app.route('/api/document/conversational-assembly', methods=['POST'])
def conversational_assembly():
    """
    SMART Conversational Document Assembly
    
    Intelligently extracts information from conversation and assembles documents.
    Uses BGE-M3 embeddings + GPT-4 for context-aware variable extraction.
    
    Request:
        {
            "user_message": "I want a rent agreement, my name is Dhruv...",
            "session_id": "session_123",
            "template_id": "Lease-Agreement",  # optional, auto-detected if not provided
            "conversation_history": [...]  # optional, for context
        }
    
    Response:
        {
            "status": "needs_more_info" | "ready_to_generate" | "generated",
            "message": "What's the monthly rent amount?",
            "extracted_variables": {"LESSEE_NAME": "Dhruv", ...},
            "missing_variables": ["MONTHLY_RENT", ...],
            "document": "..." (if status == "generated"),
            "progress": {"current": 3, "total": 8}
        }
    """
    try:
        from ai.variable_extractor import variable_extractor
        from ai.template_manager_v2 import get_template_manager
        from ai.document_assembler import document_assembler
        
        data = request.json
        user_message = data.get('user_message', '').strip()
        session_id = data.get('session_id', f"session_{uuid.uuid4()}")
        template_id = data.get('template_id')
        conversation_history = data.get('conversation_history', [])
        
        if not user_message:
            return jsonify({'error': 'user_message is required'}), 400
        
        logger.info(f"🎯 Smart Assembly | Session: {session_id} | Message: {user_message[:100]}...")
        
        # Step 1: Auto-detect template if not provided
        if not template_id:
            # Use GPT to understand what document they want
            system_prompt = """You are identifying which legal document the user needs.

Available templates:
- Lease-Agreement / Rent-Agreement (for rental properties)
- Employment-Contract / NDA  
- Legal-Notice

Return ONLY the template name (e.g., "Lease-Agreement") or "UNKNOWN" if unclear."""
            
            response = ai_service.chat_completion([
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_message}
            ], temperature=0.1, max_tokens=50)
            
            template_id = response.strip().replace('"', '').replace("'", "")
            
            if template_id == "UNKNOWN":
                return jsonify({
                    'status': 'needs_clarification',
                    'message': "What type of document do you need? (e.g., lease agreement, NDA, legal notice)",
                    'available_templates': ['Lease-Agreement', 'NDA', 'Legal-Notice']
                })
            
            logger.info(f"📋 Auto-detected template: {template_id}")
        
        # Step 2: Smart variable extraction with conversation context
        extraction_result = variable_extractor.extract_from_description(
            user_description=user_message,
            template_id=template_id,
            conversation_history=conversation_history,
            session_id=session_id
        )
        
        extracted_vars = extraction_result.get('extracted_variables', {})
        missing_vars = extraction_result.get('missing_variables', [])
        
        logger.info(f"✅ Extracted: {len(extracted_vars)} | Missing: {len(missing_vars)}")
        
        # Step 3: Clean extracted values (extract actual values from user responses)
        cleaned_vars = {}
        for var_name, var_data in extracted_vars.items():
            if isinstance(var_data, dict):
                cleaned_vars[var_name] = var_data.get('value', var_data)
            else:
                cleaned_vars[var_name] = var_data
        
        # Step 4: Check if we have enough to generate
        if not missing_vars:
            # Generate the document!
            logger.info("🎉 All variables collected, generating document...")
            
            # Load template and assemble
            template_manager_v2 = get_template_manager()
            template_doc = template_manager_v2.load_template(template_id)
            
            if not template_doc:
                return jsonify({
                    'status': 'error',
                    'message': f'Template not found: {template_id}'
                }), 404
            
            # Assemble document
            assembled_doc = document_assembler.assemble_document(
                template_doc,
                cleaned_vars,
                show_missing=True
            )
            
            # Convert to text preview
            from docx import Document
            preview_text = '\n\n'.join([p.text for p in assembled_doc.paragraphs if p.text.strip()])
            
            # Save document
            output_path = f"./generated_documents/{session_id}_{template_id}.docx"
            os.makedirs("./generated_documents", exist_ok=True)
            assembled_doc.save(output_path)
            
            return jsonify({
                'status': 'generated',
                'message': f"🎉 Your {template_id.replace('-', ' ')} is ready!",
                'document': preview_text,
                'download_url': f'/api/document/download/{session_id}',
                'extracted_variables': cleaned_vars,
                'template_id': template_id
            })
        
        else:
            # Ask for next missing variable
            prompt = variable_extractor.generate_missing_variable_prompt(
                missing_variables=missing_vars,
                template_id=template_id,
                already_provided=cleaned_vars,
                conversation_history=conversation_history
            )
            
            total_vars = len(cleaned_vars) + len(missing_vars)
            progress = {
                'current': len(cleaned_vars),
                'total': total_vars,
                'percentage': int((len(cleaned_vars) / total_vars) * 100) if total_vars > 0 else 0
            }
            
            return jsonify({
                'status': 'needs_more_info',
                'message': prompt,
                'extracted_variables': cleaned_vars,
                'missing_variables': missing_vars,
                'progress': progress,
                'template_id': template_id
            })
    
    except Exception as e:
        logger.error(f"❌ Conversational assembly error: {str(e)}", exc_info=True)
        return jsonify({
            'status': 'error',
            'message': 'Sorry, something went wrong. Please try again.',
            'error': str(e)
        }), 500


# ============================================
# USER TEMPLATE UPLOAD & CONVERSION ENDPOINTS
# ============================================

@app.route('/api/template/upload-and-analyze', methods=['POST'])
def upload_and_analyze_template():
    """
    Upload a template file and analyze placeholders
    Supports: DOCX, PDF, TXT, RTF, ODT formats
    
    Returns analysis of detected placeholders and suggested variable names
    """
    try:
        from ai.template_converter import template_converter
        from ai.document_extractor import document_extractor
        
        if 'file' not in request.files:
            return jsonify({'error': 'No file uploaded'}), 400
        
        file = request.files['file']
        
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        # Check file extension (multi-format support)
        allowed_extensions = ['.docx', '.pdf', '.txt', '.rtf', '.odt']
        file_ext = Path(file.filename).suffix.lower()
        
        if file_ext not in allowed_extensions:
            return jsonify({
                'error': f'Unsupported format: {file_ext}. Supported: {", ".join(allowed_extensions)}'
            }), 400
        
        # Save temporarily with correct extension
        import tempfile
        with tempfile.NamedTemporaryFile(delete=False, suffix=file_ext) as tmp:
            file.save(tmp.name)
            tmp_path = tmp.name
        
        logger.info(f"📤 Analyzing {document_extractor.get_file_format(tmp_path)}: {file.filename}")
        
        # Analyze template (multi-format)
        analysis = template_converter.analyze_template(tmp_path)
        
        # Store temp path for later conversion
        if analysis.get('success'):
            analysis['temp_path'] = tmp_path
            analysis['original_filename'] = file.filename
            analysis['input_format'] = file_ext.replace('.', '')
        else:
            # Clean up on failure
            os.unlink(tmp_path)
        
        return jsonify(analysis)
        
    except Exception as e:
        logger.error(f"❌ Template upload/analysis error: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/template/convert', methods=['POST'])
def convert_template():
    """
    Convert template to Jinja2 format with user-provided variable mapping
    
    Request body:
    {
        "temp_path": "/tmp/...",
        "variable_mapping": {
            "#1": "party_name_1",
            "#2": "effective_date",
            ...
        },
        "template_name": "My Custom NDA",
        "category": "Employment"
    }
    """
    try:
        from ai.template_converter import template_converter
        
        data = request.json
        temp_path = data.get('temp_path')
        variable_mapping = data.get('variable_mapping', {})
        template_name = data.get('template_name', 'Untitled Template')
        category = data.get('category', 'Custom')
        
        if not temp_path or not os.path.exists(temp_path):
            return jsonify({'error': 'Invalid temp_path'}), 400
        
        if not variable_mapping:
            return jsonify({'error': 'variable_mapping required'}), 400
        
        logger.info(f"🔄 Converting template: {template_name}")
        
        # Create output directory for user templates
        user_templates_dir = Path("data/user_templates")
        user_templates_dir.mkdir(parents=True, exist_ok=True)
        
        # Generate safe filename
        safe_name = re.sub(r'[^\w\s-]', '', template_name).strip().replace(' ', '_')
        output_filename = f"{safe_name}_jinja2.docx"
        output_path = user_templates_dir / output_filename
        
        # Convert template
        conversion_result = template_converter.convert_to_jinja2(
            temp_path,
            str(output_path),
            variable_mapping
        )
        
        if not conversion_result['success']:
            return jsonify(conversion_result), 400
        
        # Extract metadata
        metadata_result = template_converter.extract_template_metadata(
            str(output_path),
            template_name,
            category
        )
        
        if not metadata_result['success']:
            return jsonify(metadata_result), 400
        
        # Clean up temp file
        try:
            os.unlink(temp_path)
        except:
            pass
        
        return jsonify({
            'success': True,
            'conversion': conversion_result,
            'metadata': metadata_result['metadata'],
            'output_path': str(output_path),
            'message': f'Template "{template_name}" converted successfully!'
        })
        
    except Exception as e:
        logger.error(f"❌ Template conversion error: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/template/save-to-library', methods=['POST'])
def save_template_to_library():
    """
    Save converted template to library (update template_config.json)
    
    Request body:
    {
        "metadata": {...},
        "output_path": "data/user_templates/..."
    }
    """
    try:
        data = request.json
        metadata = data.get('metadata', {})
        output_path = data.get('output_path')
        
        if not metadata or not output_path:
            return jsonify({'error': 'metadata and output_path required'}), 400
        
        # Load user template config
        user_config_path = Path("data/user_templates/user_template_config.json")
        
        if user_config_path.exists():
            with open(user_config_path, 'r', encoding='utf-8') as f:
                user_config = json.load(f)
        else:
            user_config = {}
        
        # Add/update template
        template_name = metadata['name']
        user_config[template_name] = metadata
        
        # Save config
        with open(user_config_path, 'w', encoding='utf-8') as f:
            json.dump(user_config, f, indent=2, ensure_ascii=False)
        
        logger.info(f"✅ Saved template to library: {template_name}")
        
        return jsonify({
            'success': True,
            'message': f'Template "{template_name}" saved to library',
            'template_name': template_name
        })
        
    except Exception as e:
        logger.error(f"❌ Save template error: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/template/user-templates', methods=['GET'])
def get_user_templates():
    """Get list of all user-uploaded templates"""
    try:
        user_config_path = Path("data/user_templates/user_template_config.json")
        
        if not user_config_path.exists():
            return jsonify({
                'success': True,
                'templates': [],
                'count': 0
            })
        
        with open(user_config_path, 'r', encoding='utf-8') as f:
            user_config = json.load(f)
        
        # Format for frontend (skip metadata keys like _readme)
        templates = [
            {
                'name': name,
                'category': config.get('category'),
                'description': config.get('description'),
                'field_count': len(config.get('fields', {})),
                'is_user_template': True
            }
            for name, config in user_config.items()
            if isinstance(config, dict) and not name.startswith('_')
        ]
        
        return jsonify({
            'success': True,
            'templates': templates,
            'count': len(templates)
        })
        
    except Exception as e:
        logger.error(f"❌ Get user templates error: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/template/delete/<template_name>', methods=['DELETE'])
def delete_user_template(template_name):
    """Delete a user-uploaded template"""
    try:
        user_config_path = Path("data/user_templates/user_template_config.json")
        
        if not user_config_path.exists():
            return jsonify({'error': 'No user templates found'}), 404
        
        with open(user_config_path, 'r', encoding='utf-8') as f:
            user_config = json.load(f)
        
        if template_name not in user_config:
            return jsonify({'error': 'Template not found'}), 404
        
        # Get template info
        template_info = user_config[template_name]
        template_file = Path("data/user_templates") / template_info.get('filename', '')
        
        # Delete file if exists
        if template_file.exists():
            os.unlink(template_file)
            logger.info(f"🗑️ Deleted template file: {template_file}")
        
        # Remove from config
        del user_config[template_name]
        
        # Save updated config
        with open(user_config_path, 'w', encoding='utf-8') as f:
            json.dump(user_config, f, indent=2, ensure_ascii=False)
        
        logger.info(f"✅ Deleted template: {template_name}")
        
        return jsonify({
            'success': True,
            'message': f'Template "{template_name}" deleted successfully'
        })
        
    except Exception as e:
        logger.error(f"❌ Delete template error: {str(e)}")
        return jsonify({'error': str(e)}), 500


if __name__ == '__main__':
    logger.info("="*60)
    logger.info("🌐 Starting Flask Server...")
    logger.info("Server will be available at: http://127.0.0.1:5000")
    logger.info("CORS enabled for: http://localhost:3000")
    logger.info("="*60)
    logger.info("Server is ready! Waiting for requests...")
    logger.info("="*60)
    
    app.run(
        host='127.0.0.1',
        port=5000,
        debug=True,
        use_reloader=False  # Prevent double initialization
    )

