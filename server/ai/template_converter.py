"""
Template Converter Service - Multi-Format Edition
Automatically converts user-uploaded templates to Jinja2 format

Features:
- Supports DOCX, PDF, TXT, RTF, ODT formats
- GPT-powered intelligent placeholder detection
- Detects various placeholder types (hash #, underscores _, dots ., brackets [], etc.)
- Converts to Jinja2 format {{ variable_name }}
- Generates intelligent variable names using AI
- Optimized token usage for cost efficiency
- Extracts field metadata
- Validates conversion quality
"""

import re
import json
import logging
from pathlib import Path
from typing import Dict, List, Optional
from docx import Document
from .document_extractor import document_extractor
from .azure_openai_service import ai_service
from .config import AIConfig

logger = logging.getLogger(__name__)


class TemplateConverter:
    """Converts user templates to Jinja2 format - Multi-format support"""
    
    # Enhanced placeholder detection patterns
    PLACEHOLDER_PATTERNS = {
        'hash': r'#\d+',                          # #1, #2, #123
        'underscore': r'_{4,}',                   # ____, ___________
        'dots': r'\.{4,}',                        # ...., ..........
        'brackets_square': r'\[[\s\w-]+\]',       # [NAME], [DATE]
        'brackets_curly': r'\{[\s\w-]+\}',        # {NAME}, {DATE}
        'brackets_angle': r'<[\s\w-]+>',          # <NAME>, <DATE>
        'dollar': r'\$\{[\w_]+\}',                # ${NAME}, ${DATE}
        'percent': r'%[\w_]+%',                   # %NAME%, %DATE%
        'double_brackets': r'\{\{[\s\w_]+\}\}',  # {{NAME}}, {{DATE}}
    }
    
    def __init__(self):
        self.ai_enabled = AIConfig.validate()
        if not self.ai_enabled:
            logger.warning("AI service not configured - using basic variable naming")
        logger.info(f"📄 TemplateConverter initialized (AI: {self.ai_enabled})")
    
    def analyze_template(self, doc_path: str) -> Dict:
        """
        Analyze template to detect placeholders (multi-format support)
        
        Returns:
            {
                'total_placeholders': 15,
                'format': 'pdf',
                'placeholder_types': {...},
                'suggested_conversions': {...},
                'context': {...}
            }
        """
        try:
            # Check if format is supported
            if not document_extractor.is_supported(doc_path):
                file_ext = Path(doc_path).suffix
                return {
                    'success': False,
                    'error': f'Unsupported format: {file_ext}. Supported: .docx, .pdf, .txt, .rtf, .odt'
                }
            
            # Extract document content
            logger.info(f"📖 Analyzing {Path(doc_path).name}")
            extracted = document_extractor.extract(doc_path)
            
            full_text = extracted['full_text']
            doc_format = extracted['format']
            
            # Detect placeholders by type
            detected_placeholders = {}
            placeholder_contexts = {}
            
            for ptype, pattern in self.PLACEHOLDER_PATTERNS.items():
                matches = re.findall(pattern, full_text)
                if matches:
                    # Remove duplicates while preserving order
                    detected_placeholders[ptype] = list(dict.fromkeys(matches))
            
            # Get context for each placeholder
            for ptype, placeholders in detected_placeholders.items():
                for placeholder in placeholders:
                    context = self._get_placeholder_context(full_text, placeholder)
                    placeholder_contexts[placeholder] = context
            
            total = sum(len(v) for v in detected_placeholders.values())
            
            # If no placeholders detected, use GPT to find them intelligently
            if total == 0 and self.ai_enabled:
                logger.info("🤖 No regex placeholders found, using GPT-4 for intelligent detection")
                gpt_analysis = self._gpt_smart_placeholder_detection(full_text[:3000])
                if gpt_analysis.get('placeholders'):
                    detected_placeholders = gpt_analysis['placeholders']
                    placeholder_contexts = gpt_analysis.get('contexts', {})
                    total = sum(len(v) for v in detected_placeholders.values())
            
            # Generate suggested variable names using AI
            suggested_conversions = {}
            if self.ai_enabled and placeholder_contexts:
                suggested_conversions = self._generate_variable_names_ai(
                    placeholder_contexts,
                    full_text[:2000],  # Limit context to reduce tokens
                    doc_format
                )
            else:
                # Fallback to basic naming
                suggested_conversions = self._generate_variable_names_basic(detected_placeholders)
            
            return {
                'success': True,
                'format': doc_format,
                'total_placeholders': total,
                'placeholder_types': detected_placeholders,
                'suggested_conversions': suggested_conversions,
                'context': placeholder_contexts,
                'document_preview': full_text[:500],
                'metadata': extracted['metadata']
            }
            
        except Exception as e:
            logger.error(f"Template analysis failed: {e}")
            return {
                'success': False,
                'error': str(e)
            }
    
    def convert_to_jinja2(
        self,
        doc_path: str,
        output_path: str,
        variable_mapping: Dict[str, str]
    ) -> Dict:
        """
        Convert template placeholders to Jinja2 format (multi-format support)
        Always outputs as DOCX regardless of input format
        
        Args:
            doc_path: Path to original template (any supported format)
            output_path: Path to save converted template (.docx)
            variable_mapping: Dict mapping placeholders to variable names
        
        Returns:
            {
                'success': True,
                'converted_count': 15,
                'output_path': '...',
                'format': 'pdf -> docx'
            }
        """
        try:
            input_format = Path(doc_path).suffix.lower().replace('.', '')
            
            # Extract content from any format
            extracted = document_extractor.extract(doc_path)
            
            # Create new DOCX document
            doc = Document()
            converted_count = 0
            
            # Sort mappings by length (descending) to avoid partial replacements
            sorted_mappings = sorted(
                variable_mapping.items(),
                key=lambda x: len(x[0]),
                reverse=True
            )
            
            # Convert paragraphs
            for para_text in extracted['paragraphs']:
                converted_text = para_text
                for placeholder, var_name in sorted_mappings:
                    if placeholder in converted_text:
                        escaped = re.escape(placeholder)
                        new_text = re.sub(
                            escaped,
                            f"{{{{ {var_name} }}}}",
                            converted_text
                        )
                        if new_text != converted_text:
                            converted_count += 1
                            logger.info(f"✓ {placeholder} → {{{{ {var_name} }}}}")
                        converted_text = new_text
                
                doc.add_paragraph(converted_text)
            
            # Convert tables
            for table_data in extracted['tables']:
                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]) if table_data else 0)
                for i, row_data in enumerate(table_data):
                    for j, cell_text in enumerate(row_data):
                        converted_text = cell_text
                        for placeholder, var_name in sorted_mappings:
                            if placeholder in converted_text:
                                escaped = re.escape(placeholder)
                                new_text = re.sub(
                                    escaped,
                                    f"{{{{ {var_name} }}}}",
                                    converted_text
                                )
                                if new_text != converted_text:
                                    converted_count += 1
                                converted_text = new_text
                        
                        table.rows[i].cells[j].text = converted_text
            
            # Save as DOCX
            doc.save(output_path)
            
            # Check for remaining placeholders
            remaining = self._find_remaining_placeholders(output_path)
            
            logger.info(f"✅ Converted {input_format} → docx ({converted_count} placeholders)")
            
            return {
                'success': True,
                'converted_count': converted_count,
                'output_path': output_path,
                'format_conversion': f'{input_format} → docx',
                'remaining_placeholders': remaining,
                'message': f'Converted {converted_count} placeholders from {input_format} to Jinja2-ready DOCX'
            }
            
        except Exception as e:
            logger.error(f"Template conversion failed: {e}")
            return {
                'success': False,
                'error': str(e)
            }
    
    def _get_placeholder_context(self, text: str, placeholder: str, context_chars: int = 50) -> str:
        """Get surrounding text context for a placeholder"""
        try:
            index = text.find(placeholder)
            if index == -1:
                return ""
            
            start = max(0, index - context_chars)
            end = min(len(text), index + len(placeholder) + context_chars)
            
            context = text[start:end]
            # Clean up
            context = ' '.join(context.split())
            return context
        except:
            return ""
    
    def _generate_variable_names_ai(
        self,
        placeholder_contexts: Dict[str, str],
        document_preview: str,
        doc_format: str = 'docx'
    ) -> Dict[str, str]:
        """Use AI to generate meaningful variable names based on context (token-optimized)"""
        try:
            # Limit contexts to reduce token usage
            limited_contexts = dict(list(placeholder_contexts.items())[:15])
            
            # Prepare contexts for GPT
            contexts_list = [
                f"'{placeholder}': \"{context[:80]}...\""
                for placeholder, context in limited_contexts.items()
            ]
            
            # Concise prompt to save tokens
            prompt = f"""Legal template analysis. Generate variable names (snake_case) for placeholders.

Document ({doc_format}): {document_preview[:400]}...

Placeholders:
{chr(10).join(contexts_list)}

Return JSON mapping:
{{"placeholder": "variable_name"}}

Rules: snake_case, descriptive, <30 chars, legal context (party_name, date, amount, address)"""

            response = ai_service.chat_completion([
                {"role": "system", "content": "Legal doc analyzer. Return JSON only."},
                {"role": "user", "content": prompt}
            ], temperature=0.3, max_tokens=800)
            
            # Parse JSON response
            match = re.search(r'\{.*\}', response, re.DOTALL)
            if match:
                mapping = json.loads(match.group(0))
                logger.info(f"✅ AI generated {len(mapping)} variable names ({len(response)} tokens)")
                return mapping
            else:
                logger.warning("AI response not JSON, using fallback")
                return self._generate_variable_names_basic({'mixed': list(placeholder_contexts.keys())})
                
        except Exception as e:
            logger.error(f"AI variable naming failed: {e}")
            return self._generate_variable_names_basic({'mixed': list(placeholder_contexts.keys())})
    
    def _gpt_smart_placeholder_detection(self, text_sample: str) -> Dict:
        """
        Use GPT-4 to intelligently detect placeholders when regex fails
        Optimized for minimal token usage
        """
        try:
            prompt = f"""Analyze this document template and identify ALL placeholders/blank fields.

Document sample:
{text_sample}

Identify:
1. Obvious blanks: ____, ...., [NAME], {{DATE}}, etc.
2. Fields that need user input
3. Repeating patterns that should be variables

Return JSON:
{{
    "placeholders": {{
        "type": ["placeholder1", "placeholder2"]
    }},
    "contexts": {{
        "placeholder1": "context text"
    }}
}}

Keep it concise."""

            response = ai_service.chat_completion([
                {"role": "system", "content": "Template analyzer. Return JSON only."},
                {"role": "user", "content": prompt}
            ], temperature=0.2, max_tokens=1000)
            
            match = re.search(r'\{.*\}', response, re.DOTALL)
            if match:
                result = json.loads(match.group(0))
                logger.info(f"🤖 GPT detected {sum(len(v) for v in result.get('placeholders', {}).values())} placeholders")
                return result
            
            return {'placeholders': {}, 'contexts': {}}
            
        except Exception as e:
            logger.error(f"GPT smart detection failed: {e}")
            return {'placeholders': {}, 'contexts': {}}
    
    def _generate_variable_names_basic(self, detected_placeholders: Dict) -> Dict[str, str]:
        """Fallback: Generate basic variable names without AI"""
        mapping = {}
        counter = 1
        
        for ptype, placeholders in detected_placeholders.items():
            for placeholder in placeholders:
                # Try to extract meaningful name from bracket placeholders
                if ptype in ['brackets_square', 'brackets_curly', 'brackets_angle']:
                    # Extract text inside brackets
                    inner = re.sub(r'[\[\]<>{}]', '', placeholder).strip()
                    var_name = inner.lower().replace(' ', '_').replace('-', '_')
                    mapping[placeholder] = var_name
                else:
                    # Generic naming
                    mapping[placeholder] = f"field_{counter}"
                    counter += 1
        
        return mapping
    
    def _find_remaining_placeholders(self, doc_path: str) -> List[str]:
        """Check for any remaining unconverted placeholders"""
        try:
            doc = Document(doc_path)
            all_text = []
            
            for para in doc.paragraphs:
                all_text.append(para.text)
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            all_text.append(para.text)
            
            full_text = '\n'.join(all_text)
            
            # Check for remaining placeholders
            remaining = []
            for ptype, pattern in self.PLACEHOLDER_PATTERNS.items():
                matches = re.findall(pattern, full_text)
                remaining.extend(matches)
            
            return list(dict.fromkeys(remaining))  # Remove duplicates
            
        except:
            return []
    
    def validate_conversion(self, doc_path: str) -> Dict:
        """Validate that conversion was successful"""
        try:
            doc = Document(doc_path)
            
            # Count Jinja2 variables
            all_text = []
            for para in doc.paragraphs:
                all_text.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            all_text.append(para.text)
            
            full_text = '\n'.join(all_text)
            
            # Find all Jinja2 variables
            jinja_vars = re.findall(r'\{\{\s*(\w+)\s*\}\}', full_text)
            unique_vars = list(dict.fromkeys(jinja_vars))
            
            # Check for remaining placeholders
            remaining = self._find_remaining_placeholders(doc_path)
            
            return {
                'success': True,
                'jinja_variable_count': len(jinja_vars),
                'unique_variables': unique_vars,
                'remaining_placeholders': remaining,
                'is_valid': len(remaining) == 0 and len(unique_vars) > 0
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': str(e)
            }
    
    def extract_template_metadata(
        self,
        doc_path: str,
        template_name: str,
        category: str = "Custom"
    ) -> Dict:
        """Extract metadata for template configuration"""
        try:
            validation = self.validate_conversion(doc_path)
            
            if not validation['success']:
                return validation
            
            # Build field configuration
            fields = {}
            for var_name in validation['unique_variables']:
                # Convert snake_case to title case for label
                label = var_name.replace('_', ' ').title()
                
                # Detect field type based on name
                field_type = 'text'
                if any(word in var_name.lower() for word in ['date', 'day', 'month', 'year']):
                    field_type = 'date'
                elif any(word in var_name.lower() for word in ['amount', 'price', 'rent', 'fee']):
                    field_type = 'number'
                elif 'email' in var_name.lower():
                    field_type = 'email'
                
                # Mark as required by default
                fields[var_name] = {
                    'label': label,
                    'type': field_type,
                    'required': True,
                    'example': ''
                }
            
            metadata = {
                'name': template_name,
                'category': category,
                'filename': Path(doc_path).name,
                'description': f'User-uploaded template: {template_name}',
                'keywords': [template_name.lower(), category.lower()],
                'fields': fields,
                'is_user_template': True,
                'jinja_variable_count': validation['jinja_variable_count']
            }
            
            return {
                'success': True,
                'metadata': metadata
            }
            
        except Exception as e:
            logger.error(f"Metadata extraction failed: {e}")
            return {
                'success': False,
                'error': str(e)
            }


# Singleton instance
template_converter = TemplateConverter()
