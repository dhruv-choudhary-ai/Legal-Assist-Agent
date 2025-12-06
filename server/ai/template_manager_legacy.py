"""
Template Manager
Handles legal document templates with variable extraction and validation
Based on Docassemble patterns and python-docx-template
"""

import os
import re
import json
import logging
from typing import Dict, List, Optional, Set, Tuple
from pathlib import Path
from docx import Document

logger = logging.getLogger(__name__)


class TemplateManager:
    """
    Manages legal document templates with Jinja2-style variables
    
    Features:
    - Template discovery and loading
    - Variable extraction from templates
    - Variable type inference
    - Template validation
    - Metadata management
    """
    
    def __init__(self, template_dir: str = "./data/templates"):
        """
        Initialize template manager
        
        Args:
            template_dir: Directory containing template files
        """
        self.template_dir = Path(template_dir)
        self.templates_cache = {}
        self.metadata_cache = {}
        
        # Ensure template directory exists
        self.template_dir.mkdir(parents=True, exist_ok=True)
        
        # Variable pattern matching
        # Supports: {{variable}}, {variable}, [VARIABLE], [[variable]]
        self.variable_patterns = [
            r'\{\{([^}]+)\}\}',  # Jinja2 style {{variable}}
            r'\{([A-Z_][A-Z0-9_]*)\}',  # {VARIABLE}
            r'\[([A-Z_][A-Z0-9_\s]*)\]',  # [VARIABLE NAME]
            r'\[\[([^\]]+)\]\]'  # [[variable]]
        ]
        
        logger.info(f"📂 Template Manager initialized | Directory: {self.template_dir}")
    
    def discover_templates(self) -> Dict[str, Dict]:
        """
        Discover all templates in the template directory
        
        Returns:
            Dict mapping template_id to template info
        """
        templates = {}
        
        for category_dir in self.template_dir.iterdir():
            if not category_dir.is_dir():
                continue
            
            category = category_dir.name
            
            for template_file in category_dir.glob("*.docx"):
                if template_file.name.startswith("~$"):  # Skip temp files
                    continue
                
                template_id = f"{category}/{template_file.stem}"
                
                # Extract variables from template
                variables = self.extract_variables(template_id)
                
                templates[template_id] = {
                    'id': template_id,
                    'name': template_file.stem.replace('_', ' ').title(),
                    'category': category,
                    'file_path': str(template_file),
                    'file_name': template_file.name,
                    'variable_count': len(variables),
                    'variables': list(variables.keys())
                }
        
        logger.info(f"🔍 Discovered {len(templates)} templates across {len(set(t['category'] for t in templates.values()))} categories")
        return templates
    
    def load_template(self, template_id: str) -> Optional[Document]:
        """
        Load a template document
        
        Args:
            template_id: Template identifier (e.g., "employment/nda")
        
        Returns:
            python-docx Document object or None
        """
        # Check cache
        if template_id in self.templates_cache:
            logger.info(f"📦 Loading template from cache: {template_id}")
            return self.templates_cache[template_id]
        
        # Find template file
        template_path = self.template_dir / f"{template_id}.docx"
        
        if not template_path.exists():
            logger.error(f"❌ Template not found: {template_id}")
            return None
        
        try:
            doc = Document(template_path)
            self.templates_cache[template_id] = doc
            logger.info(f"✅ Template loaded: {template_id}")
            return doc
        
        except Exception as e:
            logger.error(f"❌ Failed to load template {template_id}: {e}")
            return None
    
    def extract_text_from_doc(self, doc: Document) -> str:
        """
        Extract all text from document (paragraphs and tables)
        
        Args:
            doc: python-docx Document
        
        Returns:
            Full document text
        """
        text_parts = []
        
        # Extract from paragraphs
        for para in doc.paragraphs:
            text_parts.append(para.text)
        
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_parts.append(cell.text)
        
        return "\n".join(text_parts)
    
    def extract_variables(self, template_id: str) -> Dict[str, Dict]:
        """
        Extract all variables from a template
        
        Args:
            template_id: Template identifier
        
        Returns:
            Dict mapping variable name to variable info:
            {
                "PARTY_NAME": {
                    "name": "PARTY_NAME",
                    "display_name": "Party Name",
                    "type": "text",
                    "required": True,
                    "description": "Name of the party",
                    "default": None,
                    "example": "John Doe"
                }
            }
        """
        doc = self.load_template(template_id)
        if not doc:
            return {}
        
        # Extract text
        text = self.extract_text_from_doc(doc)
        
        # Find all variables
        variables_found = set()
        
        for pattern in self.variable_patterns:
            matches = re.finditer(pattern, text)
            for match in matches:
                var_name = match.group(1).strip().upper().replace(' ', '_')
                variables_found.add(var_name)
        
        # Build variable info
        variables = {}
        
        for var_name in sorted(variables_found):
            var_info = self._infer_variable_info(var_name, text)
            variables[var_name] = var_info
        
        logger.info(f"🔤 Extracted {len(variables)} variables from {template_id}")
        return variables
    
    def _infer_variable_info(self, var_name: str, context_text: str) -> Dict:
        """
        Infer variable type and metadata from its name and context
        
        Args:
            var_name: Variable name (e.g., "PARTY_NAME")
            context_text: Surrounding text for context
        
        Returns:
            Variable info dict
        """
        # Clean display name
        display_name = var_name.replace('_', ' ').title()
        
        # Infer type from name patterns
        var_type = "text"  # default
        
        if any(keyword in var_name.lower() for keyword in ['date', 'day', 'month', 'year']):
            var_type = "date"
        elif any(keyword in var_name.lower() for keyword in ['amount', 'salary', 'price', 'fee', 'cost', 'rupees']):
            var_type = "currency"
        elif any(keyword in var_name.lower() for keyword in ['email', 'mail']):
            var_type = "email"
        elif any(keyword in var_name.lower() for keyword in ['phone', 'mobile', 'contact']):
            var_type = "phone"
        elif any(keyword in var_name.lower() for keyword in ['address', 'location', 'city', 'state']):
            var_type = "address"
        elif any(keyword in var_name.lower() for keyword in ['number', 'count', 'quantity']):
            var_type = "number"
        elif any(keyword in var_name.lower() for keyword in ['description', 'details', 'clause']):
            var_type = "textarea"
        
        # Generate description and example
        description = self._generate_description(var_name)
        example = self._generate_example(var_name, var_type)
        
        return {
            'name': var_name,
            'display_name': display_name,
            'type': var_type,
            'required': True,  # All template variables are required by default
            'description': description,
            'default': None,
            'example': example
        }
    
    def _generate_description(self, var_name: str) -> str:
        """Generate human-readable description for variable"""
        parts = var_name.lower().split('_')
        
        # Common prefixes
        if parts[0] == 'party':
            return f"Details of {'first' if '1' in var_name else 'second' if '2' in var_name else 'the'} party"
        elif parts[0] == 'employee':
            return f"Employee's {' '.join(parts[1:])}"
        elif parts[0] == 'employer':
            return f"Employer's {' '.join(parts[1:])}"
        elif parts[0] == 'agreement':
            return f"Agreement {' '.join(parts[1:])}"
        
        return f"Provide {var_name.replace('_', ' ').lower()}"
    
    def _generate_example(self, var_name: str, var_type: str) -> str:
        """Generate example value for variable"""
        examples = {
            'text': {
                'name': 'John Doe',
                'company': 'ABC Pvt Ltd',
                'title': 'Software Engineer',
                'designation': 'Senior Developer',
                'pan': 'ABCDE1234F',
                'aadhar': '1234 5678 9012'
            },
            'date': '2025-01-01',
            'currency': '₹800,000',
            'email': 'john.doe@example.com',
            'phone': '+91 98765 43210',
            'address': '123 MG Road, Bangalore, Karnataka 560001',
            'number': '5',
            'textarea': 'Detailed description here...'
        }
        
        # Try to find specific example
        for key, value in examples.get('text', {}).items():
            if key in var_name.lower():
                return value
        
        # Return type-based example
        return examples.get(var_type, 'Example value')
    
    def get_template_metadata(self, template_id: str) -> Dict:
        """
        Get complete metadata for a template
        
        Args:
            template_id: Template identifier
        
        Returns:
            Complete template metadata including variables
        """
        # Get basic info
        templates = self.discover_templates()
        template_info = templates.get(template_id)
        
        if not template_info:
            return {}
        
        # Get variables
        variables = self.extract_variables(template_id)
        
        # Load document for stats
        doc = self.load_template(template_id)
        
        metadata = {
            **template_info,
            'variables': variables,
            'variable_count': len(variables),
            'required_variables': [v for v, info in variables.items() if info['required']],
            'optional_variables': [v for v, info in variables.items() if not info['required']],
            'statistics': {
                'paragraphs': len(doc.paragraphs) if doc else 0,
                'tables': len(doc.tables) if doc else 0,
                'sections': len(doc.sections) if doc else 0
            }
        }
        
        return metadata
    
    def validate_template(self, template_id: str) -> Tuple[bool, List[str]]:
        """
        Validate template structure and variables
        
        Args:
            template_id: Template identifier
        
        Returns:
            Tuple of (is_valid, list of errors/warnings)
        """
        errors = []
        
        # Check if template exists
        doc = self.load_template(template_id)
        if not doc:
            errors.append(f"Template file not found: {template_id}")
            return False, errors
        
        # Extract variables
        variables = self.extract_variables(template_id)
        
        if not variables:
            errors.append("Warning: Template has no variables - may be a static document")
        
        # Check for common issues
        text = self.extract_text_from_doc(doc)
        
        # Check for mismatched brackets
        if text.count('{{') != text.count('}}'):
            errors.append("Mismatched {{ }} brackets in template")
        
        if text.count('[[') != text.count(']]'):
            errors.append("Mismatched [[ ]] brackets in template")
        
        # Check for minimum content
        if len(text) < 100:
            errors.append("Warning: Template seems too short")
        
        # Check for essential legal sections (basic heuristic)
        essential_keywords = ['whereas', 'agreement', 'party', 'terms', 'conditions']
        found_keywords = sum(1 for kw in essential_keywords if kw.lower() in text.lower())
        
        if found_keywords < 2:
            errors.append("Warning: Template may be missing essential legal sections")
        
        is_valid = all('Error' not in e for e in errors)
        
        logger.info(f"✓ Template validation: {template_id} - {'VALID' if is_valid else 'INVALID'} ({len(errors)} issues)")
        return is_valid, errors
    
    def create_template_index(self) -> Dict:
        """
        Create searchable index of all templates
        
        Returns:
            Template index with metadata
        """
        templates = self.discover_templates()
        index = {}
        
        for template_id, template_info in templates.items():
            metadata = self.get_template_metadata(template_id)
            index[template_id] = metadata
        
        # Save index to file
        index_path = self.template_dir / "template_index.json"
        with open(index_path, 'w', encoding='utf-8') as f:
            json.dump(index, f, indent=2, ensure_ascii=False)
        
        logger.info(f"📑 Template index created: {len(index)} templates indexed")
        return index


# Global template manager instance
template_manager = TemplateManager()
