"""
Convert all templates to Jinja2 format
Reads existing templates and creates Jinja2 versions with proper variable mapping
"""

import json
from pathlib import Path
from docx import Document
import re

# Path setup
BASE_DIR = Path(__file__).parent.parent
TEMPLATES_DIR = BASE_DIR / "data" / "templates"
PLACEHOLDER_FILE = TEMPLATES_DIR / "placeholder_details.json"
CONFIG_FILE = TEMPLATES_DIR / "template_config.json"

# Load placeholder analysis
with open(PLACEHOLDER_FILE, 'r', encoding='utf-8') as f:
    placeholder_data = json.load(f)

# Load existing config
with open(CONFIG_FILE, 'r', encoding='utf-8') as f:
    config = json.load(f)

print("="*80)
print("JINJA2 TEMPLATE CONVERSION TOOL")
print("="*80)

# Template conversion mappings
CONVERSIONS = {
    "Deed of Lease .docx": {
        "new_filename": "Lease-Agreement-Hash-to-Jinja2.docx",
        "placeholders": {
            "#1": "city",
            "#2": "day",
            "#3": "month",
            "#18": "year",
            "#4": "lessor_name",
            "#5": "lessor_address",
            "#6": "lessee_name",
            "#7": "lessee_address",
            "#8": "lease_duration_years",
            "#9": "property_address",
            "#10": "start_month",
            "#11": "start_year",
            "#12": "monthly_rent",
            "#13": "interest_rate",
            "#15": "notice_period_months",
            "#16": "lessor_witness_1",
            "#17": "lessee_witness_2"
        },
        "config_name": "Lease Agreement (Alternative)"
    },
    
    "Legal-Notice-for-Recovery-of-Money.docx": {
        "new_filename": "Legal-Notice-for-Recovery-of-Money-Jinja2.docx",
        "placeholders": {
            # Map underscore placeholders by their context/position
            # This is tricky - we'll use a sequential approach
        },
        "sequential_fields": [
            "notice_date_day",
            "notice_date_month",
            "notice_date_year", 
            "recipient_name",
            "recipient_address",
            "client_name",
            "client_designation",
            "client_full_name",
            "business_nature",
            "business_product",
            "principal_amount",
            "interest_rate",
            "interest_amount",
            "total_payable",
            "notice_fee",
            "advocate_name"
        ],
        "config_name": "Legal Notice for Recovery of Money"
    },
    
    "Deed-of-Family-Trust.docx": {
        "new_filename": "Family-Trust-Deed-Jinja2.docx",
        "sequential_fields": [
            "trust_location",
            "trust_day",
            "trust_month",
            "trust_year",
            "settlor_1_name",
            "settlor_2_name", 
            "settlor_3_name",
            "settlor_4_name",
            "settlor_5_name",
            "settlor_6_name",
            "settlor_7_name",
            "settlor_city",
            "trustee_1_name",
            "trustee_2_name",
            "trustee_3_name",
            "brother_1_name",
            "brother_2_name",
            "brother_3_name",
            "property_village",
            "property_district",
            "property_state",
            "deity_name",
            "initial_fund_amount",
            "trust_name",
            "property_address"
        ],
        "config_name": "Family Trust Deed"
    }
}

def convert_hash_placeholders(doc_path, output_path, placeholder_map):
    """Convert hash placeholders (#1, #2) to Jinja2 variables"""
    doc = Document(doc_path)
    
    # Sort placeholders by length (descending) to avoid partial replacements
    sorted_placeholders = sorted(placeholder_map.items(), key=lambda x: len(x[0]), reverse=True)
    
    # Replace in paragraphs
    for para in doc.paragraphs:
        for placeholder, var_name in sorted_placeholders:
            if placeholder in para.text:
                # Replace placeholder with Jinja2 variable
                para.text = para.text.replace(placeholder, f"{{{{ {var_name} }}}}")
                print(f"  ✓ Replaced {placeholder} → {{{{ {var_name} }}}}")
    
    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for placeholder, var_name in sorted_placeholders:
                        if placeholder in para.text:
                            para.text = para.text.replace(placeholder, f"{{{{ {var_name} }}}}")
                            print(f"  ✓ Replaced {placeholder} → {{{{ {var_name} }}}} (in table)")
    
    doc.save(output_path)
    print(f"✅ Saved: {output_path}")

def convert_underscore_placeholders(doc_path, output_path, field_list):
    """Convert underscore placeholders to Jinja2 variables"""
    doc = Document(doc_path)
    
    field_index = 0
    
    # Replace in paragraphs
    for para in doc.paragraphs:
        text = para.text
        # Find all underscore sequences
        underscore_matches = re.findall(r'_{4,}', text)
        
        for match in underscore_matches:
            if field_index < len(field_list):
                var_name = field_list[field_index]
                text = text.replace(match, f"{{{{ {var_name} }}}}", 1)  # Replace only first occurrence
                print(f"  ✓ Replaced {match} → {{{{ {var_name} }}}}")
                field_index += 1
        
        para.text = text
    
    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text
                    underscore_matches = re.findall(r'_{4,}', text)
                    
                    for match in underscore_matches:
                        if field_index < len(field_list):
                            var_name = field_list[field_index]
                            text = text.replace(match, f"{{{{ {var_name} }}}}", 1)
                            print(f"  ✓ Replaced {match} → {{{{ {var_name} }}}} (in table)")
                            field_index += 1
                    
                    para.text = text
    
    doc.save(output_path)
    print(f"✅ Saved: {output_path}")

def convert_dot_placeholders(doc_path, output_path, field_list):
    """Convert dot placeholders to Jinja2 variables"""
    doc = Document(doc_path)
    
    field_index = 0
    
    # Replace in paragraphs
    for para in doc.paragraphs:
        text = para.text
        # Find all dot sequences (10+ dots)
        dot_matches = re.findall(r'\.{10,}', text)
        
        for match in dot_matches:
            if field_index < len(field_list):
                var_name = field_list[field_index]
                text = text.replace(match, f"{{{{ {var_name} }}}}", 1)
                print(f"  ✓ Replaced {match} → {{{{ {var_name} }}}}")
                field_index += 1
        
        para.text = text
    
    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text
                    dot_matches = re.findall(r'\.{10,}', text)
                    
                    for match in dot_matches:
                        if field_index < len(field_list):
                            var_name = field_list[field_index]
                            text = text.replace(match, f"{{{{ {var_name} }}}}", 1)
                            print(f"  ✓ Replaced {match} → {{{{ {var_name} }}}} (in table)")
                            field_index += 1
                    
                    para.text = text
    
    doc.save(output_path)
    print(f"✅ Saved: {output_path}")

# Process each template
print("\n📁 Processing templates...\n")

# 1. Convert Deed of Lease (hash placeholders)
if (TEMPLATES_DIR / "Deed of Lease .docx").exists():
    print("🔄 Converting: Deed of Lease .docx")
    convert_hash_placeholders(
        TEMPLATES_DIR / "Deed of Lease .docx",
        TEMPLATES_DIR / "Lease-Agreement-Hash-to-Jinja2.docx",
        CONVERSIONS["Deed of Lease .docx"]["placeholders"]
    )
    print()

# 2. Convert Legal Notices (underscore placeholders)
for notice_file in ["Legal-Notice-for-Recovery-of-Money.docx", 
                     "Legal-Notice-for-Non-Payment-of-Invoice.docx",
                     "Legal-Notice-for-Recovery-of-Friendly-Loan.docx"]:
    if (TEMPLATES_DIR / notice_file).exists():
        print(f"🔄 Converting: {notice_file}")
        output_name = notice_file.replace(".docx", "-Jinja2.docx")
        convert_underscore_placeholders(
            TEMPLATES_DIR / notice_file,
            TEMPLATES_DIR / output_name,
            CONVERSIONS["Legal-Notice-for-Recovery-of-Money.docx"]["sequential_fields"]
        )
        print()

# 3. Convert Family Trust (dot placeholders)
if (TEMPLATES_DIR / "Deed-of-Family-Trust.docx").exists():
    print("🔄 Converting: Deed-of-Family-Trust.docx")
    convert_dot_placeholders(
        TEMPLATES_DIR / "Deed-of-Family-Trust.docx",
        TEMPLATES_DIR / "Family-Trust-Deed-Jinja2.docx",
        CONVERSIONS["Deed-of-Family-Trust.docx"]["sequential_fields"]
    )
    print()

print("\n" + "="*80)
print("✅ CONVERSION COMPLETE!")
print("="*80)
print("\nNext steps:")
print("1. Review the new Jinja2 templates in the templates folder")
print("2. Update template_config.json with the new filenames")
print("3. Test each template by creating a document")
print("\nNew templates created:")
print("  - Lease-Agreement-Hash-to-Jinja2.docx")
print("  - Legal-Notice-for-Recovery-of-Money-Jinja2.docx")
print("  - Legal-Notice-for-Non-Payment-of-Invoice-Jinja2.docx")
print("  - Legal-Notice-for-Recovery-of-Friendly-Loan-Jinja2.docx")
print("  - Family-Trust-Deed-Jinja2.docx")
