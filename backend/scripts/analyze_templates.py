"""
Analyze existing .docx templates to extract placeholder variables
"""
import os
from docx import Document
from docxtpl import DocxTemplate
import re
import json

def analyze_docx_template(filepath):
    """Extract all placeholders from a .docx template"""
    try:
        # Use python-docx to read the document
        doc = Document(filepath)
        
        # Get all text from document
        full_text = ""
        for paragraph in doc.paragraphs:
            full_text += paragraph.text + "\n"
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text += cell.text + "\n"
        
        # Find Jinja2-style placeholders: {{ variable }}
        jinja_placeholders = re.findall(r'\{\{\s*(\w+)\s*\}\}', full_text)
        
        # Find bracket-style placeholders: [VARIABLE]
        bracket_placeholders = re.findall(r'\[([A-Z_\s]+)\]', full_text)
        
        # Find underscore placeholders: _________
        underscore_placeholders = re.findall(r'_{3,}', full_text)
        
        return {
            'jinja2_variables': list(set(jinja_placeholders)),
            'bracket_variables': list(set(bracket_placeholders)),
            'underscore_count': len(underscore_placeholders),
            'sample_text': full_text[:500] if full_text else "No text found"
        }
    except Exception as e:
        return {'error': str(e)}

def main():
    templates_dir = os.path.join(os.path.dirname(__file__), '..', 'data', 'templates')
    
    # Get all .docx files
    template_files = [
        'Deed of Lease .docx',
        'Deed-of-Family-Trust.docx',
        'Lease-Deed-(for-a-term-of-years)-Rent-Agreement.docx',
        'Legal-Notice-for-Non-Payment-of-Invoice.docx',
        'Legal-Notice-for-Recovery-of-Friendly-Loan.docx',
        'Legal-Notice-for-Recovery-of-Money.docx'
    ]
    
    results = {}
    
    print("=" * 80)
    print("TEMPLATE ANALYSIS REPORT")
    print("=" * 80)
    
    for template_file in template_files:
        filepath = os.path.join(templates_dir, template_file)
        
        if not os.path.exists(filepath):
            print(f"\n❌ NOT FOUND: {template_file}")
            continue
        
        print(f"\n📄 Template: {template_file}")
        print("-" * 80)
        
        analysis = analyze_docx_template(filepath)
        results[template_file] = analysis
        
        if 'error' in analysis:
            print(f"   ❌ Error: {analysis['error']}")
        else:
            print(f"   Jinja2 Variables ({{ var }}): {len(analysis['jinja2_variables'])}")
            if analysis['jinja2_variables']:
                for var in analysis['jinja2_variables'][:10]:  # Show first 10
                    print(f"      • {var}")
            
            print(f"\n   Bracket Variables [VAR]: {len(analysis['bracket_variables'])}")
            if analysis['bracket_variables']:
                for var in analysis['bracket_variables'][:10]:
                    print(f"      • {var}")
            
            print(f"\n   Underscores (___): {analysis['underscore_count']}")
            
            print(f"\n   Sample Text:")
            print(f"   {analysis['sample_text'][:300]}...")
    
    # Save results to JSON
    output_file = os.path.join(templates_dir, 'template_analysis.json')
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(results, f, indent=2)
    
    print("\n" + "=" * 80)
    print(f"✅ Analysis saved to: {output_file}")
    print("=" * 80)

if __name__ == '__main__':
    main()
