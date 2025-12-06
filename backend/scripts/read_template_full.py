"""
Read full content of a specific template
"""
from docx import Document
import sys

def read_template(filepath):
    doc = Document(filepath)
    
    print("=" * 100)
    print("FULL TEMPLATE CONTENT")
    print("=" * 100)
    
    for i, paragraph in enumerate(doc.paragraphs, 1):
        if paragraph.text.strip():
            print(f"\n[Para {i}] {paragraph.text}")
    
    # Check for tables
    if doc.tables:
        print("\n" + "=" * 100)
        print("TABLES")
        print("=" * 100)
        for i, table in enumerate(doc.tables, 1):
            print(f"\n[Table {i}]")
            for row in table.rows:
                row_text = " | ".join([cell.text for cell in row.cells])
                print(f"  {row_text}")

if __name__ == '__main__':
    import os
    filepath = os.path.join(os.path.dirname(__file__), '..', 'data', 'templates', 
                            'Lease-Deed-(for-a-term-of-years)-Rent-Agreement.docx')
    read_template(filepath)
