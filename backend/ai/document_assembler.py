"""
Document Assembler
Assembles legal documents by filling templates with variables
Based on python-docx-template patterns
"""

import logging
import re
from typing import Dict, Optional
from pathlib import Path
from docx import Document
from copy import deepcopy

logger = logging.getLogger(__name__)


class DocumentAssembler:
    """
    Assembles legal documents from templates and variables
    
    Features:
    - Variable substitution in paragraphs and tables
    - Preserve document formatting
    - Generate preview with placeholders
    - Export to DOCX
    """
    
    def __init__(self):
        logger.info("📝 Document Assembler initialized")
    
    def assemble_document(
        self,
        template_doc: Document,
        variables: Dict[str, str],
        show_missing: bool = True
    ) -> Document:
        """
        Assemble document by replacing variables with values
        
        Args:
            template_doc: Template Document object
            variables: Dict mapping variable names to values
            show_missing: Show [MISSING: VAR] for unfilled variables
        
        Returns:
            Assembled Document
        """
        # Create copy of template
        assembled_doc = deepcopy(template_doc)
        
        # Replace in paragraphs
        for para in assembled_doc.paragraphs:
            self._replace_in_paragraph(para, variables, show_missing)
        
        # Replace in tables
        for table in assembled_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self._replace_in_paragraph(para, variables, show_missing)
        
        logger.info(f"✅ Document assembled with {len(variables)} variables")
        return assembled_doc
    
    def _replace_in_paragraph(
        self,
        paragraph,
        variables: Dict[str, str],
        show_missing: bool
    ):
        """Replace variables in a paragraph while preserving formatting"""
        
        # Patterns to match
        patterns = [
            (r'\{\{([^}]+)\}\}', '{{', '}}'),  # {{VAR}}
            (r'\{([A-Z_][A-Z0-9_]*)\}', '{', '}'),  # {VAR}
            (r'\[([A-Z_][A-Z0-9_\s]*)\]', '[', ']'),  # [VAR]
            (r'\[\[([^\]]+)\]\]', '[[', ']]')  # [[VAR]]
        ]
        
        text = paragraph.text
        
        for pattern, open_br, close_br in patterns:
            matches = list(re.finditer(pattern, text))
            
            for match in reversed(matches):  # Reverse to maintain positions
                var_name = match.group(1).strip().upper().replace(' ', '_')
                
                if var_name in variables:
                    replacement = variables[var_name]
                else:
                    replacement = f"[MISSING: {var_name}]" if show_missing else f"{open_br}{var_name}{close_br}"
                
                # Replace in text
                start, end = match.span()
                text = text[:start] + replacement + text[end:]
        
        # Update paragraph text (this will lose some formatting)
        # For production, use python-docx-template for better formatting preservation
        if text != paragraph.text:
            paragraph.text = text
    
    def generate_preview(
        self,
        template_doc: Document,
        variables: Dict[str, str],
        missing_variables: list
    ) -> str:
        """
        Generate text preview of document
        
        Args:
            template_doc: Template document
            variables: Available variables
            missing_variables: List of missing variable names
        
        Returns:
            Text preview with variable status
        """
        preview_parts = []
        preview_parts.append("=== DOCUMENT PREVIEW ===\n")
        
        # Show filled variables
        if variables:
            preview_parts.append("✅ FILLED VARIABLES:")
            for var, value in variables.items():
                preview_parts.append(f"  • {var}: {value}")
            preview_parts.append("")
        
        # Show missing variables
        if missing_variables:
            preview_parts.append("⚠️  MISSING VARIABLES:")
            for var in missing_variables:
                preview_parts.append(f"  • {var}: [TO BE PROVIDED]")
            preview_parts.append("")
        
        # Generate partial document text
        assembled = self.assemble_document(template_doc, variables, show_missing=True)
        
        preview_parts.append("📄 DOCUMENT CONTENT (First 500 chars):")
        preview_parts.append("-" * 50)
        
        text_parts = []
        for para in assembled.paragraphs[:10]:  # First 10 paragraphs
            if para.text.strip():
                text_parts.append(para.text)
        
        preview_text = "\n\n".join(text_parts)
        preview_parts.append(preview_text[:500] + "..." if len(preview_text) > 500 else preview_text)
        
        return "\n".join(preview_parts)
    
    def export_document(
        self,
        assembled_doc: Document,
        output_path: str
    ) -> bool:
        """
        Export assembled document to file
        
        Args:
            assembled_doc: Assembled Document
            output_path: Output file path
        
        Returns:
            Success status
        """
        try:
            output_file = Path(output_path)
            output_file.parent.mkdir(parents=True, exist_ok=True)
            
            assembled_doc.save(str(output_file))
            logger.info(f"💾 Document exported: {output_path}")
            return True
        
        except Exception as e:
            logger.error(f"❌ Export failed: {e}")
            return False
    
    def validate_assembly(
        self,
        assembled_doc: Document
    ) -> Dict:
        """
        Validate assembled document
        
        Args:
            assembled_doc: Assembled document
        
        Returns:
            Validation results
        """
        results = {
            'is_complete': True,
            'missing_variables': [],
            'warnings': []
        }
        
        # Extract all text
        text_parts = []
        for para in assembled_doc.paragraphs:
            text_parts.append(para.text)
        
        for table in assembled_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        text_parts.append(para.text)
        
        full_text = "\n".join(text_parts)
        
        # Check for missing variables
        missing_pattern = r'\[MISSING:\s*([^\]]+)\]'
        missing_matches = re.findall(missing_pattern, full_text)
        
        if missing_matches:
            results['is_complete'] = False
            results['missing_variables'] = list(set(missing_matches))
        
        # Check for unfilled placeholders
        placeholder_patterns = [
            r'\{\{[^}]+\}\}',
            r'\{[A-Z_][A-Z0-9_]*\}',
            r'\[([A-Z_][A-Z0-9_\s]*)\]'
        ]
        
        for pattern in placeholder_patterns:
            unfilled = re.findall(pattern, full_text)
            if unfilled:
                results['warnings'].append(f"Unfilled placeholders found: {unfilled[:3]}")
        
        return results


# Global instance
document_assembler = DocumentAssembler()
