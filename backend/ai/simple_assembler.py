"""
Simple Document Assembler - No Chaos, Just Works
User prompt → GPT extracts → Fill template → Done
"""

import logging
import json
import re
from typing import Dict, List, Optional
from pathlib import Path
from docx import Document
from .azure_openai_service import ai_service
from .rag_pipeline import rag_pipeline

logger = logging.getLogger(__name__)

# Template configurations
TEMPLATE_CONFIG = {
    "Lease Agreement": {
        "file": "Deed of Lease .docx",
        "fields": {
            "#1": "CITY",
            "#2": "DAY", 
            "#3": "MONTH",
            "#18": "YEAR",
            "#4": "LESSOR_NAME",
            "#5": "LESSOR_ADDRESS",
            "#6": "LESSEE_NAME",
            "#7": "LESSEE_ADDRESS",
            "#8": "LEASE_DURATION_YEARS",
            "#9": "PROPERTY_ADDRESS",
            "#10": "START_MONTH",
            "#11": "START_YEAR",
            "#12": "MONTHLY_RENT",
            "#13": "INTEREST_RATE",
            "#15": "NOTICE_PERIOD_MONTHS",
            "#16": "WITNESS_1_NAME",
            "#17": "WITNESS_2_NAME"
        }
    },
    "NDA": {
        "file": "nda_template.docx",
        "fields": {
            "PARTY_1_NAME": "First party name",
            "PARTY_2_NAME": "Second party name",
            "EFFECTIVE_DATE": "Agreement date",
            "JURISDICTION": "Legal jurisdiction"
        }
    },
    "Legal Notice": {
        "file": "Legal-Notice-for-Recovery-of-Money.docx",
        "fields": {
            "____": "Date",
            "__________": "Notice number/Recipient",
            "_____________": "Recipient name",
            "___________________": "Client name"
        }
    }
}


class SimpleAssembler:
    """Dead simple document assembler"""
    
    def __init__(self):
        self.template_dir = Path("data/templates")
        logger.info("✅ Simple Assembler initialized")
    
    def detect_template(self, user_prompt: str) -> Optional[str]:
        """Let GPT pick the template from our list"""
        
        template_list = list(TEMPLATE_CONFIG.keys())
        
        prompt = f"""Which document does the user want? Pick ONE from this list:
{json.dumps(template_list)}

User said: "{user_prompt}"

Reply with ONLY the exact template name from the list, or "UNKNOWN"."""

        response = ai_service.chat_completion([
            {"role": "system", "content": "You are a document type classifier. Reply with only the template name."},
            {"role": "user", "content": prompt}
        ], temperature=0.1, max_tokens=50)
        
        detected = response.strip().strip('"').strip("'")
        
        if detected in template_list:
            logger.info(f"✅ Detected: {detected}")
            return detected
        
        logger.warning(f"⚠️ Unknown template: {detected}")
        return None
    
    def extract_fields(self, user_prompt: str, template_name: str, conversation: List[Dict] = None) -> Dict:
        """GPT extracts ALL fields from prompt + conversation"""
        
        config = TEMPLATE_CONFIG.get(template_name, {})
        required_fields = config.get("fields", {})
        
        # Build context from conversation
        context = ""
        if conversation:
            context = "PREVIOUS CONVERSATION:\n"
            for msg in conversation[-5:]:
                context += f"{msg['role']}: {msg['content']}\n"
        
        prompt = f"""{context}

CURRENT USER MESSAGE: "{user_prompt}"

Extract these fields (CLEAN VALUES ONLY - no phrases like "I told you" or "my name is"):

{json.dumps(required_fields, indent=2)}

CRITICAL EXTRACTION RULES:
1. Extract ACTUAL values: "owner is Rahul" → extract "Rahul", NOT "owner is Rahul"
2. "my name is X" → extract "X" 
3. "I told you Y" → extract "Y"
4. "company TechVita" → extract "TechVita"
5. For #4 (LESSOR_NAME) and #6 (LESSEE_NAME):
   - If user says "my company" or "for my company" → that's the LESSEE (tenant)
   - If user says "with X" or "owner X" → that's the LESSOR (landlord)
6. For dates: extract day, month name, year separately
7. For money: just the number (e.g., "20k" → "20000")

Examples:
- "rent agreement for TechVita with Mahesh Kumar" → LESSEE_NAME: "TechVita", LESSOR_NAME: "Mahesh Kumar"
- "I told you Mahesh Kumar" → LESSOR_NAME: "Mahesh Kumar"
- "lease is 20k for 3 years" → MONTHLY_RENT: "20000", LEASE_DURATION_YEARS: "3"

Return JSON (ONLY include fields you found):
{{
  "#4": "landlord name here",
  "#6": "tenant name here",
  "#12": "rent amount here",
  ...
}}

Use the field codes (#1, #4, #6, etc.) as keys, not the descriptions."""

        response = ai_service.chat_completion([
            {"role": "system", "content": "You extract field values cleanly. Return JSON with field codes as keys."},
            {"role": "user", "content": prompt}
        ], temperature=0.1, max_tokens=1000)
        
        # Parse JSON
        try:
            match = re.search(r'\{.*\}', response, re.DOTALL)
            if match:
                extracted = json.loads(match.group(0))
                logger.info(f"📥 Raw extraction: {extracted}")
            else:
                extracted = {}
        except Exception as e:
            logger.error(f"Failed to parse: {response}")
            extracted = {}
        
        # Map placeholder codes to field names for storage
        mapped_extracted = {}
        for placeholder, value in extracted.items():
            field_name = required_fields.get(placeholder)
            if field_name:
                mapped_extracted[field_name] = value
                logger.info(f"✅ Mapped {placeholder} ({field_name}) = {value}")
        
        # Calculate what's missing (by placeholder codes)
        extracted_placeholders = set(extracted.keys())
        required_placeholders = set(required_fields.keys())
        missing_placeholders = list(required_placeholders - extracted_placeholders)
        
        logger.info(f"✅ Extracted {len(extracted)}/{len(required_fields)} fields | Missing: {missing_placeholders[:5]}")
        
        return {
            "extracted": mapped_extracted,  # Field names → values
            "extracted_raw": extracted,  # Placeholder codes → values
            "missing": missing_placeholders,  # Placeholder codes
            "required": required_fields  # All fields
        }
    
    def ask_for_missing(self, missing_fields: List[str], template_name: str, already_have_raw: Dict) -> str:
        """Generate natural question for missing field - ONLY ask for truly missing fields
        
        Args:
            missing_fields: List of missing placeholder codes (e.g., ["#4", "#6"])
            template_name: Template name
            already_have_raw: Dict of placeholder codes we already have (e.g., {"#4": "Mahesh", "#6": "TechVita"})
        """
        
        if not missing_fields:
            return "Perfect! I have everything I need."
        
        config = TEMPLATE_CONFIG.get(template_name, {})
        field_info = config.get("fields", {})
        
        # Filter out what we ACTUALLY already have
        truly_missing = []
        for placeholder in missing_fields:
            # If we already have this placeholder code, skip it
            if placeholder in already_have_raw:
                logger.info(f"✅ Skipping {placeholder} ({field_info.get(placeholder)}) - already have: {already_have_raw[placeholder]}")
            else:
                truly_missing.append(placeholder)
        
        if not truly_missing:
            logger.info("✅ No truly missing fields after checking!")
            return "Perfect! I have all the information I need. Generating document..."
        
        # Pick first truly missing field
        next_placeholder = truly_missing[0]
        field_desc = field_info.get(next_placeholder, next_placeholder)
        
        # Show what we have (convert to readable format)
        have_summary = ', '.join([
            f"{field_info.get(k, k)}: {v}" 
            for k, v in list(already_have_raw.items())[:3]
        ])
        
        prompt = f"""Ask for this ONE field naturally (one short sentence):

Need: {field_desc}
Already have: {have_summary}

Be conversational. Don't ask for info we already have."""

        question = ai_service.chat_completion([
            {"role": "system", "content": "Ask for ONE missing piece of document info. Be brief and natural."},
            {"role": "user", "content": prompt}
        ], temperature=0.7, max_tokens=50)
        
        return question.strip()
    
    def fill_template(self, template_name: str, fields: Dict) -> Document:
        """Fill template with field values (fields should be placeholder code → value)"""
        
        config = TEMPLATE_CONFIG.get(template_name)
        if not config:
            raise ValueError(f"Unknown template: {template_name}")
        
        template_file = self.template_dir / config["file"]
        if not template_file.exists():
            raise FileNotFoundError(f"Template not found: {template_file}")
        
        doc = Document(template_file)
        
        logger.info(f"🔧 Filling template with {len(fields)} values")
        
        # Replace placeholders in all paragraphs and tables
        for para in doc.paragraphs:
            original_text = para.text
            for placeholder, value in fields.items():
                if placeholder in para.text:
                    para.text = para.text.replace(placeholder, str(value))
                    logger.info(f"  Replaced {placeholder} with '{value}' in paragraph")
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for placeholder, value in fields.items():
                            if placeholder in para.text:
                                para.text = para.text.replace(placeholder, str(value))
                                logger.info(f"  Replaced {placeholder} with '{value}' in table")
        
        logger.info(f"✅ Template filled: {template_name}")
        return doc
    
    def enhance_with_rag(self, document_text: str, template_name: str) -> str:
        """Use BGE-M3 RAG to suggest additional legal clauses"""
        
        try:
            query = f"Legal clauses and provisions for {template_name}"
            results = rag_pipeline.search_knowledge_base(query, n_results=3)
            
            if results:
                suggestions = "\n\nSUGGESTED ADDITIONAL CLAUSES (from legal database):\n"
                for i, result in enumerate(results, 1):
                    suggestions += f"{i}. {result.get('text', '')[:200]}...\n"
                return suggestions
            
        except Exception as e:
            logger.warning(f"RAG enhancement failed: {e}")
        
        return ""


# Singleton
simple_assembler = SimpleAssembler()
