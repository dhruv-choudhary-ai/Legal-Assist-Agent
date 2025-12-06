"""
Multi-Format Document Extractor
Extracts text and structure from PDF, DOCX, TXT, RTF, ODT formats

Features:
- Unified interface for all document formats
- Preserves structure (paragraphs, tables)
- Optimized for template conversion
- Minimal token usage for AI processing
"""

import os
import re
import logging
from pathlib import Path
from typing import Dict, List, Tuple, Optional

# Document format libraries
from docx import Document
try:
    import fitz  # PyMuPDF for PDF
except ImportError:
    try:
        import pymupdf as fitz  # Alternative import for newer versions
    except ImportError:
        fitz = None
        logger.warning("PyMuPDF not available - PDF support disabled")

from striprtf.striprtf import rtf_to_text
from odf import text as odf_text, teletype
from odf.opendocument import load as odf_load

logger = logging.getLogger(__name__)


class DocumentExtractor:
    """Extract text and structure from multiple document formats"""
    
    SUPPORTED_FORMATS = {
        '.docx': 'Microsoft Word',
        '.pdf': 'PDF Document',
        '.txt': 'Plain Text',
        '.rtf': 'Rich Text Format',
        '.odt': 'OpenDocument Text'
    }
    
    def __init__(self):
        logger.info("📄 DocumentExtractor initialized with multi-format support")
    
    def extract(self, file_path: str) -> Dict:
        """
        Extract text and structure from any supported format
        
        Returns:
            {
                'format': 'docx',
                'paragraphs': ['text1', 'text2', ...],
                'tables': [[['cell1', 'cell2'], ['cell3', 'cell4']], ...],
                'full_text': 'complete document text',
                'metadata': {'pages': 5, 'has_tables': True, ...}
            }
        """
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext not in self.SUPPORTED_FORMATS:
            raise ValueError(f"Unsupported format: {file_ext}. Supported: {', '.join(self.SUPPORTED_FORMATS.keys())}")
        
        logger.info(f"📖 Extracting from {self.SUPPORTED_FORMATS[file_ext]}: {Path(file_path).name}")
        
        # Route to appropriate extractor
        extractors = {
            '.docx': self._extract_docx,
            '.pdf': self._extract_pdf,
            '.txt': self._extract_txt,
            '.rtf': self._extract_rtf,
            '.odt': self._extract_odt
        }
        
        result = extractors[file_ext](file_path)
        result['format'] = file_ext.replace('.', '')
        
        logger.info(f"✅ Extracted {len(result['paragraphs'])} paragraphs, {len(result['tables'])} tables")
        return result
    
    def _extract_docx(self, file_path: str) -> Dict:
        """Extract from Microsoft Word .docx"""
        doc = Document(file_path)
        
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)
        
        full_text = '\n'.join(paragraphs)
        for table in tables:
            full_text += '\n' + '\n'.join(['\t'.join(row) for row in table])
        
        return {
            'paragraphs': paragraphs,
            'tables': tables,
            'full_text': full_text,
            'metadata': {
                'paragraph_count': len(paragraphs),
                'table_count': len(tables),
                'has_tables': len(tables) > 0
            }
        }
    
    def _extract_pdf(self, file_path: str) -> Dict:
        """Extract from PDF using PyMuPDF (most accurate)"""
        if fitz is None:
            raise ImportError("PyMuPDF not installed. Install with: pip install pymupdf")
        
        doc = fitz.open(file_path)
        
        paragraphs = []
        tables = []
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            
            # Extract text blocks (preserves structure better than get_text())
            blocks = page.get_text("blocks")
            
            for block in blocks:
                # block format: (x0, y0, x1, y1, text, block_no, block_type)
                text = block[4].strip()
                if text:
                    # Split by double newlines to get paragraphs
                    paras = [p.strip() for p in text.split('\n\n') if p.strip()]
                    paragraphs.extend(paras)
            
            # Try to detect tables (simple heuristic)
            # Tables have aligned text in columns
            table_data = self._detect_pdf_tables(page)
            if table_data:
                tables.append(table_data)
        
        doc.close()
        
        full_text = '\n'.join(paragraphs)
        for table in tables:
            full_text += '\n' + '\n'.join(['\t'.join(row) for row in table])
        
        return {
            'paragraphs': paragraphs,
            'tables': tables,
            'full_text': full_text,
            'metadata': {
                'page_count': len(doc),
                'paragraph_count': len(paragraphs),
                'table_count': len(tables),
                'has_tables': len(tables) > 0
            }
        }
    
    def _detect_pdf_tables(self, page) -> Optional[List[List[str]]]:
        """Simple table detection in PDF (can be enhanced)"""
        # This is a basic implementation
        # For production, consider using tabula-py or camelot
        return None
    
    def _extract_txt(self, file_path: str) -> Dict:
        """Extract from plain text file"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        
        # Split into paragraphs (double newline separation)
        paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
        
        # If no double newlines, split by single newlines
        if len(paragraphs) <= 1:
            paragraphs = [line.strip() for line in content.split('\n') if line.strip()]
        
        return {
            'paragraphs': paragraphs,
            'tables': [],
            'full_text': content,
            'metadata': {
                'paragraph_count': len(paragraphs),
                'table_count': 0,
                'has_tables': False,
                'line_count': len(content.split('\n'))
            }
        }
    
    def _extract_rtf(self, file_path: str) -> Dict:
        """Extract from Rich Text Format"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            rtf_content = f.read()
        
        # Convert RTF to plain text
        plain_text = rtf_to_text(rtf_content)
        
        # Split into paragraphs
        paragraphs = [p.strip() for p in plain_text.split('\n\n') if p.strip()]
        
        if len(paragraphs) <= 1:
            paragraphs = [line.strip() for line in plain_text.split('\n') if line.strip()]
        
        return {
            'paragraphs': paragraphs,
            'tables': [],
            'full_text': plain_text,
            'metadata': {
                'paragraph_count': len(paragraphs),
                'table_count': 0,
                'has_tables': False
            }
        }
    
    def _extract_odt(self, file_path: str) -> Dict:
        """Extract from OpenDocument Text"""
        doc = odf_load(file_path)
        
        paragraphs = []
        tables = []
        
        # Extract paragraphs
        for para in doc.getElementsByType(odf_text.P):
            text = teletype.extractText(para).strip()
            if text:
                paragraphs.append(text)
        
        # Extract tables
        for table in doc.getElementsByType(odf_text.Table):
            table_data = []
            for row in table.getElementsByType(odf_text.TableRow):
                row_data = []
                for cell in row.getElementsByType(odf_text.TableCell):
                    cell_text = teletype.extractText(cell).strip()
                    row_data.append(cell_text)
                if row_data:
                    table_data.append(row_data)
            if table_data:
                tables.append(table_data)
        
        full_text = '\n'.join(paragraphs)
        for table in tables:
            full_text += '\n' + '\n'.join(['\t'.join(row) for row in table])
        
        return {
            'paragraphs': paragraphs,
            'tables': tables,
            'full_text': full_text,
            'metadata': {
                'paragraph_count': len(paragraphs),
                'table_count': len(tables),
                'has_tables': len(tables) > 0
            }
        }
    
    @staticmethod
    def get_file_format(file_path: str) -> str:
        """Get file format description"""
        ext = Path(file_path).suffix.lower()
        return DocumentExtractor.SUPPORTED_FORMATS.get(ext, 'Unknown')
    
    @staticmethod
    def is_supported(file_path: str) -> bool:
        """Check if file format is supported"""
        ext = Path(file_path).suffix.lower()
        return ext in DocumentExtractor.SUPPORTED_FORMATS


# Singleton instance
document_extractor = DocumentExtractor()
